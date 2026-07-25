from __future__ import annotations

import argparse
import json
from hashlib import sha256
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
WARNING_FILL = "FFF4CE"
ALLOWED_REVIEW_STATUSES = {
    "transcribed_from_user_supplied_sources",
    "technically_reviewed",
}


def _require_non_blank(payload: dict[str, Any], key: str) -> str:
    value = payload.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{key} must be a non-blank string")
    return value.strip()


def _json_text(value: Any) -> str:
    if isinstance(value, str):
        return value
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _set_east_asia_font(run: Any, font_name: str) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(
    style: Any,
    *,
    size: float,
    color: RGBColor | None = None,
) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    _set_style_font(title, size=18, color=DARK_BLUE)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)

    heading_1 = document.styles["Heading 1"]
    _set_style_font(heading_1, size=16, color=BLUE)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)

    heading_2 = document.styles["Heading 2"]
    _set_style_font(heading_2, size=13, color=BLUE)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)


def _add_source_note(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"来源定位：{text}")
    _set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _validate_fact(raw_fact: Any, index: int) -> dict[str, Any]:
    if not isinstance(raw_fact, dict):
        raise ValueError(f"facts[{index}] must be an object")
    field = _require_non_blank(raw_fact, "field").lower()
    label = _require_non_blank(raw_fact, "label")
    value_type = _require_non_blank(raw_fact, "value_type")
    source_reference = _require_non_blank(raw_fact, "source_reference")
    if "value" not in raw_fact:
        raise ValueError(f"facts[{index}].value is required")
    return {
        "field": field,
        "label": label,
        "value": raw_fact["value"],
        "value_type": value_type,
        "source_reference": source_reference,
    }


def build_fact_sheet(
    *,
    input_path: Path,
    output_path: Path,
    locators_path: Path | None = None,
) -> tuple[Path, Path]:
    payload = json.loads(input_path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("fact sheet input must be a JSON object")

    case_id = _require_non_blank(payload, "case_id")
    title = _require_non_blank(payload, "title")
    review_status = _require_non_blank(payload, "review_status")
    if review_status not in ALLOWED_REVIEW_STATUSES:
        raise ValueError(f"unsupported review_status: {review_status}")
    document_version_id = payload.get("document_version_id")
    if not isinstance(document_version_id, int) or document_version_id <= 0:
        raise ValueError("document_version_id must be a positive integer")

    raw_facts = payload.get("facts")
    if not isinstance(raw_facts, list) or not raw_facts:
        raise ValueError("facts must be a non-empty list")
    facts = [_validate_fact(raw_fact, index) for index, raw_fact in enumerate(raw_facts)]
    fields = [fact["field"] for fact in facts]
    if len(fields) != len(set(fields)):
        raise ValueError("fact fields must be unique within one fact sheet")

    document = Document()
    _configure_document(document)
    title_paragraph = document.add_paragraph(title, style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    control = document.add_paragraph()
    control.paragraph_format.left_indent = Inches(0.12)
    control.paragraph_format.right_indent = Inches(0.12)
    control.paragraph_format.space_before = Pt(4)
    control.paragraph_format.space_after = Pt(10)
    _shade_paragraph(control, WARNING_FILL)
    control_run = control.add_run(
        "仅限本地离线评测。事实由开发人员从用户提供资料转录；"
        "除非 review_status=technically_reviewed，否则不代表技术负责人已审核。"
    )
    _set_east_asia_font(control_run, "Microsoft YaHei")
    control_run.bold = True

    document.add_heading("已核对事实", level=1)
    locator_facts: list[dict[str, Any]] = []
    for fact in facts:
        fact_text = f"[{fact['field']}] {fact['label']}：{_json_text(fact['value'])}"
        paragraph_index = len(document.paragraphs)
        paragraph = document.add_paragraph()
        key_run = paragraph.add_run(f"[{fact['field']}] {fact['label']}：")
        _set_east_asia_font(key_run, "Microsoft YaHei")
        key_run.bold = True
        value_run = paragraph.add_run(_json_text(fact["value"]))
        _set_east_asia_font(value_run, "Microsoft YaHei")
        _add_source_note(document, fact["source_reference"])
        locator_facts.append(
            {
                "field": fact["field"],
                "value": fact["value"],
                "value_type": fact["value_type"],
                "source_document_version_id": document_version_id,
                "locator": {
                    "paragraph_index": paragraph_index,
                    "text_quote": fact_text[:200],
                },
                "confidence": 1.0,
                "confirmed_by": 1,
            }
        )

    document.add_heading("控制信息", level=1)
    document.add_paragraph(f"盲测案例：{case_id}")
    document.add_paragraph(f"事实底稿版本号：{document_version_id}")
    document.add_paragraph(f"复核状态：{review_status}")
    document.add_paragraph("confirmed_by=1 仅为离线评测准备人占位，不映射生产用户。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    output_hash = sha256(output_path.read_bytes()).hexdigest()
    if locators_path is None:
        locators_path = output_path.with_suffix(".locators.json")
    locators_path.parent.mkdir(parents=True, exist_ok=True)
    locators_path.write_text(
        json.dumps(
            {
                "case_id": case_id,
                "document_version_id": document_version_id,
                "review_status": review_status,
                "fact_sheet_path": str(output_path.resolve()),
                "fact_sheet_sha256": output_hash,
                "confirmed_facts": locator_facts,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return output_path, locators_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Build a private, source-located Phase 5 fact sheet.",
    )
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--locators-output", type=Path)
    args = parser.parse_args(argv)
    output_path, locators_path = build_fact_sheet(
        input_path=args.input.resolve(),
        output_path=args.output.resolve(),
        locators_path=args.locators_output.resolve() if args.locators_output else None,
    )
    print(f"[PASS] fact sheet: {output_path}")
    print(f"[PASS] locators: {locators_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
