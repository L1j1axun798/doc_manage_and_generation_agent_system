from __future__ import annotations

import json
from hashlib import sha256
from io import BytesIO

import pytest
from docx import Document

from apps.document_generation.engine.contracts import KnowledgeChunk
from apps.document_generation.engine.fakes import HashingEmbeddingProvider
from scripts.document_agent.phase5_cli import run_offline


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_phase5_fake_cli_writes_docx_trace_and_validation(tmp_path) -> None:
    source_path = tmp_path / "current-project-source.docx"
    source_path.write_bytes(
        _docx_bytes(
            "入场任务资料",
            "计划开展当前风电机组检测项目。",
        )
    )
    template_path = tmp_path / "approved-template.docx"
    template_path.write_bytes(_docx_bytes("四措两案模板"))
    query_text = "overview project_name 当前风电机组检测项目 risk_evidence_items []"
    embedding_provider = HashingEmbeddingProvider()
    knowledge_text = query_text
    knowledge_chunk = KnowledgeChunk(
        chunk_id="dv134:overview:test",
        source_document_version_id=134,
        business_type="wind_turbine_inspection_four_measures_two_plans",
        section_code="overview",
        heading_path=("工程概况",),
        paragraph_start=1,
        paragraph_end=1,
        text=knowledge_text,
        approval_status="approved",
        content_sha256=sha256(knowledge_text.encode()).hexdigest(),
        embedding=tuple(embedding_provider.embed((knowledge_text,))[0]),
        embedding_model_alias=embedding_provider.model_alias,
        embedding_dimension=embedding_provider.dimension,
    )
    knowledge_path = tmp_path / "knowledge.json"
    knowledge_path.write_text(
        json.dumps([knowledge_chunk.model_dump(mode="json")], ensure_ascii=False),
        encoding="utf-8",
    )
    input_path = tmp_path / "input.json"
    input_path.write_text(
        json.dumps(
            {
                "request_id": "offline-test-001",
                "idempotency_key": "offline-test-001",
                "business_type": "wind_turbine_inspection_four_measures_two_plans",
                "template_id": "T-TEST",
                "template_path": str(template_path),
                "sources": [
                    {
                        "document_version_id": 1001,
                        "path": str(source_path),
                    }
                ],
                "confirmed_facts": [
                    {
                        "field": "project_name",
                        "value": "当前风电机组检测项目",
                        "value_type": "string",
                        "source_document_version_id": 1001,
                        "locator": {"paragraph_index": 1},
                        "confidence": 1,
                        "confirmed_by": 1,
                    },
                    {
                        "field": "risk_evidence_items",
                        "value": [],
                        "value_type": "list[object]",
                        "source_document_version_id": 1001,
                        "locator": {"paragraph_index": 1},
                        "confidence": 1,
                        "confirmed_by": 1,
                    },
                    {
                        "field": "inspection_quantity",
                        "value": 19,
                        "value_type": "integer",
                        "source_document_version_id": 1001,
                        "locator": {"paragraph_index": 1},
                        "confidence": 1,
                        "confirmed_by": 1,
                    },
                    {
                        "field": "inspection_unit",
                        "value": "台",
                        "value_type": "string",
                        "source_document_version_id": 1001,
                        "locator": {"paragraph_index": 1},
                        "confidence": 1,
                        "confirmed_by": 1,
                    },
                ],
                "required_fact_fields": [
                    "project_name",
                    "risk_evidence_items",
                    "inspection_quantity",
                    "inspection_unit",
                ],
                "section_codes": ["overview"],
                "knowledge_json_path": str(knowledge_path),
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    output_dir = tmp_path / "output"

    docx_path, trace_path, validation_path = run_offline(
        input_path=input_path,
        output_dir=output_dir,
        provider_mode="fake",
    )

    assert docx_path.is_file()
    assert trace_path.is_file()
    assert validation_path.is_file()
    review_bundle_path = output_dir / "review_bundle.json"
    assert review_bundle_path.is_file()
    document = Document(docx_path)
    assert "当前风电机组检测项目" in "\n".join(paragraph.text for paragraph in document.paragraphs)
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    validation = json.loads(validation_path.read_text(encoding="utf-8"))
    review_bundle = json.loads(review_bundle_path.read_text(encoding="utf-8"))
    assert trace["document_purpose"] == "entry_four_measures_two_plans"
    assert validation["valid"] is True
    assert validation["fact_citation_coverage"] == 1.0
    assert review_bundle["schema_version"] == "phase5-review-v1"
    assert len(review_bundle["implementation_fingerprint"]) == 64
    assert review_bundle["sections"][0]["section_code"] == "overview"
    assert review_bundle["sections"][0]["retrieved_references"][0]["chunk_id"] == (
        "dv134:overview:test"
    )


def test_phase5_cli_rejects_blind_answer_before_reading_source(tmp_path) -> None:
    input_path = tmp_path / "input.json"
    input_path.write_text(
        json.dumps(
            {
                "request_id": "blind-leak-test",
                "idempotency_key": "blind-leak-test",
                "business_type": "wind_turbine_inspection_four_measures_two_plans",
                "template_id": "T-TEST",
                "template_path": str(tmp_path / "missing-template.docx"),
                "sources": [
                    {
                        "document_version_id": 180,
                        "path": str(tmp_path / "missing-blind-answer.docx"),
                    }
                ],
                "confirmed_facts": [],
                "required_fact_fields": [],
                "section_codes": ["overview"],
            }
        ),
        encoding="utf-8",
    )
    output_dir = tmp_path / "output"

    with pytest.raises(ValueError, match="blind answer documents cannot be used"):
        run_offline(
            input_path=input_path,
            output_dir=output_dir,
            provider_mode="fake",
        )

    assert not output_dir.exists()
