from __future__ import annotations

import re
from collections.abc import Mapping
from copy import deepcopy
from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate
from jinja2 import Environment, StrictUndefined

from .contracts import (
    GeneratedSection,
    RenderImageAsset,
    RenderedArtifact,
    RenderRequest,
    TemplateDocument,
    TemplateValidationResult,
)
from .docx_formatting import (
    ParagraphFormatSample,
    TemplateFormatProfile,
    apply_body_paragraph_layout,
    apply_paragraph_format,
    apply_table_format,
    clean_generated_text,
    ensure_page_number_fields,
    format_list_item,
    infer_template_format_profile,
    is_subheading_text,
    matching_template_table,
    paragraph_sample,
    without_outline_level,
)
from .errors import SourcePurposeMismatchError, TemplateInvalidError
from .parsing import REPORT_FILENAME_MARKERS

PLACEHOLDER_RE = re.compile(r"{{\s*([A-Za-z_][A-Za-z0-9_.]*)\s*}}")
STYLE_ONLY_TEMPLATE_WARNING = "模板不含占位符，将仅继承样式并清除历史正文和页眉页脚"
SECTION_NUMBER_PREFIXES = ("一", "二", "三", "四", "五", "六", "七", "八")
CHAPTER_HEADING_RE = re.compile(
    r"^第\s*[一二三四五六七八九十百0-9]+\s*章\s*[、,.，．]?\s*(?P<title>.+?)\s*$"
)
SECTION_TITLE_ALIASES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("organization_measures", ("组织措施", "组织机构", "组织保障")),
    ("construction_plan", ("施工方案", "作业方案", "实施方案")),
    ("technical_measures", ("技术措施", "技术方案")),
    ("safety_measures", ("安全措施", "安全保障")),
    ("risk_identification", ("风险辨识", "风险预控", "危险点", "风险管控")),
    ("emergency_plan", ("应急预案", "现场处置方案", "应急处置")),
    ("environmental_measures", ("环保措施", "环境保护", "文明施工")),
    ("overview", ("项目简介", "项目概况", "工程概况", "编制依据")),
)


@dataclass(frozen=True)
class TemplateOutlineItem:
    paragraph_index: int
    section_code: str
    title: str


def _section_code_for_heading(text: str) -> str | None:
    match = CHAPTER_HEADING_RE.match(" ".join(text.split()))
    if match is None:
        return None
    title = match.group("title")
    for section_code, aliases in SECTION_TITLE_ALIASES:
        if any(alias in title for alias in aliases):
            return section_code
    return None


def template_outline(document: DocumentObject) -> tuple[TemplateOutlineItem, ...]:
    outline: list[TemplateOutlineItem] = []
    seen: set[str] = set()
    for index, paragraph in enumerate(document.paragraphs):
        section_code = _section_code_for_heading(paragraph.text)
        if section_code is None or section_code in seen:
            continue
        seen.add(section_code)
        outline.append(
            TemplateOutlineItem(
                paragraph_index=index,
                section_code=section_code,
                title=" ".join(paragraph.text.split()),
            )
        )
    return tuple(outline)


def infer_template_section_order(content: bytes) -> tuple[str, ...]:
    try:
        document = Document(BytesIO(content))
    except Exception:
        return ()
    return tuple(item.section_code for item in template_outline(document))


def _docx_xml_text(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            return "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            )
    except BadZipFile as exc:
        raise TemplateInvalidError("模板不是有效的DOCX文件") from exc


def _ensure_entry_template(template: TemplateDocument) -> None:
    if any(marker in template.filename for marker in REPORT_FILENAME_MARKERS):
        raise SourcePurposeMismatchError("报告模板不能用于入场四措两案")


def _prepare_style_only_baseline(
    source_document: DocumentObject,
    context: Mapping[str, object],
    profile: TemplateFormatProfile,
) -> DocumentObject:
    document = Document()
    document.part._styles_part._element = deepcopy(source_document.part._styles_part._element)
    document.part.numbering_part._element = deepcopy(source_document.part.numbering_part._element)

    # Historical baselines may end with a landscape appendix. The first section is
    # the approved cover/body baseline; copying the last section made every newly
    # generated entry plan landscape.
    source_section = source_document.sections[0]
    target_section = document.sections[0]
    for attribute in (
        "page_height",
        "page_width",
        "orientation",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "gutter",
        "header_distance",
        "footer_distance",
    ):
        setattr(target_section, attribute, getattr(source_section, attribute))

    project_name = context.get("project_name")
    if not isinstance(project_name, str) or not project_name.strip():
        raise TemplateInvalidError("样式基线渲染必须提供项目名称")
    _add_cover_line(document, project_name.strip(), profile.cover)
    _add_cover_line(document, "入场四措两案", profile.cover)
    client_name = context.get("client_name")
    if isinstance(client_name, str) and client_name.strip():
        _add_cover_line(
            document,
            f"委托方：{client_name.strip()}",
            profile.body,
            centered=True,
        )
    document.add_page_break()
    return document


def _add_cover_line(
    document: DocumentObject,
    text: str,
    sample: ParagraphFormatSample,
    *,
    centered: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    apply_paragraph_format(paragraph, sample)
    if centered:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _has_style(document: DocumentObject, style_name: str) -> bool:
    return any(style.name == style_name for style in document.styles)


def _add_generated_table(
    document: DocumentObject,
    *,
    rows: int,
    cols: int,
) -> Table:
    table = document.add_table(rows=rows, cols=cols)
    if _has_style(document, "Table Grid"):
        table.style = "Table Grid"
    return table


def _render_into_template_outline(
    document: DocumentObject,
    sections: tuple[GeneratedSection, ...],
    outline: tuple[TemplateOutlineItem, ...],
    profile: TemplateFormatProfile,
    image_assets: Mapping[str, RenderImageAsset],
    prototype_tables: tuple[Table, ...],
) -> None:
    section_by_code = {section.section_code: section for section in sections}
    paragraphs = list(document.paragraphs)
    heading_paragraphs = [paragraphs[item.paragraph_index] for item in outline]
    heading_paragraphs[0].paragraph_format.page_break_before = True
    body = document.element.body

    for outline_index in range(len(outline) - 1, -1, -1):
        item = outline[outline_index]
        section = section_by_code.get(item.section_code)
        if section is None:
            continue
        heading = heading_paragraphs[outline_index]
        next_heading = (
            heading_paragraphs[outline_index + 1]
            if outline_index + 1 < len(heading_paragraphs)
            else None
        )
        removable: list[object] = []
        body_examples: list[Paragraph] = []
        subheading_example: Paragraph | None = None
        source_tables: list[Table] = []
        sibling = heading._p.getnext()
        while sibling is not None and (next_heading is None or sibling is not next_heading._p):
            if sibling.tag == qn("w:sectPr"):
                break
            if sibling.tag == qn("w:p"):
                paragraph = Paragraph(sibling, document)
                text = paragraph.text.strip()
                if text and is_subheading_text(text):
                    if subheading_example is None:
                        subheading_example = paragraph
                elif text:
                    body_examples.append(paragraph)
            elif sibling.tag == qn("w:tbl"):
                source_tables.append(Table(sibling, document))
            removable.append(sibling)
            sibling = sibling.getnext()
        for element in removable:
            body.remove(element)

        anchor = heading._p
        body_sample = without_outline_level(
            paragraph_sample(
                next((value for value in body_examples if len(value.text.strip()) >= 20), None)
                or (body_examples[0] if body_examples else None)
            )
        )
        if body_sample.paragraph_properties is None and body_sample.run_properties is None:
            body_sample = profile.body
        subheading_sample = (
            paragraph_sample(subheading_example)
            if subheading_example is not None
            else profile.subheading
        )
        for raw_text in section.paragraphs:
            text = clean_generated_text(raw_text)
            is_subheading = is_subheading_text(text)
            paragraph = document.add_paragraph(text)
            apply_paragraph_format(
                paragraph,
                subheading_sample if is_subheading else body_sample,
            )
            if not is_subheading:
                apply_body_paragraph_layout(paragraph)
            anchor.addnext(paragraph._p)
            anchor = paragraph._p
        for items in section.lists:
            for item_index, value in enumerate(items, start=1):
                paragraph = document.add_paragraph(
                    format_list_item(value, index=item_index, profile=profile)
                )
                apply_paragraph_format(paragraph, profile.list_item)
                anchor.addnext(paragraph._p)
                anchor = paragraph._p
        for generated_table in section.tables:
            table = _add_generated_table(
                document,
                rows=1,
                cols=len(generated_table.headers),
            )
            for index, header in enumerate(generated_table.headers):
                table.rows[0].cells[index].text = clean_generated_text(header)
            for row in generated_table.rows:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = clean_generated_text(value)
            apply_table_format(
                (
                    prototype_tables[generated_table.prototype_table_index]
                    if generated_table.prototype_table_index is not None
                    and generated_table.prototype_table_index < len(prototype_tables)
                    else matching_template_table(
                        source_tables,
                        columns=len(generated_table.headers),
                    )
                ),
                table,
                profile=profile,
            )
            anchor.addnext(table._tbl)
            anchor = table._tbl
        for generated_image in section.images:
            asset = image_assets.get(generated_image.asset_id)
            if asset is None:
                raise TemplateInvalidError(
                    f"结构化图片缺少已确认文件：{generated_image.block_key}"
                )
            image_paragraph, caption_paragraph = _add_generated_image(
                document,
                asset=asset,
                caption=generated_image.caption,
                alt_text=generated_image.alt_text,
                profile=profile,
            )
            anchor.addnext(image_paragraph._p)
            image_paragraph._p.addnext(caption_paragraph._p)
            anchor = caption_paragraph._p


def _add_generated_image(
    document: DocumentObject,
    *,
    asset: RenderImageAsset,
    caption: str,
    alt_text: str,
    profile: TemplateFormatProfile,
) -> tuple[Paragraph, Paragraph]:
    if sha256(asset.content).hexdigest() != asset.sha256:
        raise TemplateInvalidError("结构化图片哈希校验失败")
    section = document.sections[-1]
    available_inches = float(
        (section.page_width - section.left_margin - section.right_margin) / 914400
    )
    target_inches = min(available_inches, asset.width_px / 200)
    if target_inches < 2:
        raise TemplateInvalidError("结构化图片在200 DPI下过小，无法保证A4阅读清晰度")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    run.add_picture(BytesIO(asset.content), width=Inches(target_inches))
    for drawing_property in image_paragraph._p.xpath(".//wp:docPr"):
        drawing_property.set("descr", alt_text)
        drawing_property.set("title", alt_text)
    caption_paragraph = document.add_paragraph(clean_generated_text(caption))
    apply_paragraph_format(caption_paragraph, profile.body)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_together = True
    return image_paragraph, caption_paragraph


class DocxTemplateRenderer:
    _media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def preflight(
        self,
        template: TemplateDocument,
        available_context: Mapping[str, object] | None = None,
    ) -> TemplateValidationResult:
        _ensure_entry_template(template)
        if Path(template.filename).suffix.lower() != ".docx":
            return TemplateValidationResult(
                valid=False,
                missing_placeholders=template.required_placeholders,
                warnings=("模板必须为DOCX格式",),
            )
        try:
            docx_template = DocxTemplate(BytesIO(template.content))
            declared = set(docx_template.get_undeclared_template_variables())
        except Exception:
            return TemplateValidationResult(
                valid=False,
                missing_placeholders=template.required_placeholders,
                warnings=("模板无法解析",),
            )
        xml_text = _docx_xml_text(template.content)
        declared.update(PLACEHOLDER_RE.findall(xml_text))
        required = set(template.required_placeholders)
        missing_in_template = required - declared
        context_fields = set((available_context or {}).keys())
        missing_context = declared - context_fields if available_context is not None else set()
        missing = missing_in_template | missing_context
        return TemplateValidationResult(
            valid=not missing,
            declared_placeholders=tuple(sorted(declared)),
            missing_placeholders=tuple(sorted(missing)),
            warnings=(STYLE_ONLY_TEMPLATE_WARNING,) if not declared else (),
        )

    def render(self, request: RenderRequest) -> RenderedArtifact:
        context = {fact.field: fact.value for fact in request.facts}
        validation = self.preflight(request.template, context)
        if not validation.valid:
            missing = "、".join(validation.missing_placeholders)
            raise TemplateInvalidError(f"模板预检失败，缺少字段或占位符：{missing}")

        try:
            source_document = Document(BytesIO(request.template.content))
            outline = template_outline(source_document)
            profile = infer_template_format_profile(
                source_document,
                chapter_indexes={item.paragraph_index for item in outline},
            )
            if not validation.declared_placeholders:
                document = _prepare_style_only_baseline(source_document, context, profile)
            else:
                template = DocxTemplate(BytesIO(request.template.content))
                environment = Environment(
                    autoescape=True,
                    undefined=StrictUndefined,
                )
                template.render(context, jinja_env=environment)
                rendered_buffer = BytesIO()
                template.save(rendered_buffer)
                rendered_buffer.seek(0)
                document = Document(rendered_buffer)
        except Exception as exc:
            if isinstance(exc, TemplateInvalidError):
                raise
            raise TemplateInvalidError("模板渲染失败") from exc

        image_assets = {asset.asset_id: asset for asset in request.image_assets}
        sections_to_append = tuple(request.sections)

        for section_index, section in enumerate(sections_to_append):
            prefix = (
                SECTION_NUMBER_PREFIXES[section_index]
                if section_index < len(SECTION_NUMBER_PREFIXES)
                else str(section_index + 1)
            )
            heading = document.add_paragraph(
                f"{prefix}、{clean_generated_text(section.title)}"
            )
            apply_paragraph_format(heading, profile.chapter_heading)
            for raw_paragraph in section.paragraphs:
                paragraph = clean_generated_text(raw_paragraph)
                is_subheading = is_subheading_text(paragraph)
                rendered_paragraph = document.add_paragraph(paragraph)
                apply_paragraph_format(
                    rendered_paragraph,
                    profile.subheading if is_subheading else profile.body,
                )
                if not is_subheading:
                    apply_body_paragraph_layout(rendered_paragraph)
            for items in section.lists:
                for item_index, item in enumerate(items, start=1):
                    rendered_item = document.add_paragraph(
                        format_list_item(item, index=item_index, profile=profile)
                    )
                    apply_paragraph_format(rendered_item, profile.list_item)
            for generated_table in section.tables:
                table = _add_generated_table(
                    document,
                    rows=1,
                    cols=len(generated_table.headers),
                )
                for index, header in enumerate(generated_table.headers):
                    table.rows[0].cells[index].text = clean_generated_text(header)
                for row in generated_table.rows:
                    cells = table.add_row().cells
                    for index, value in enumerate(row):
                        cells[index].text = clean_generated_text(value)
                apply_table_format(
                    (
                        source_document.tables[generated_table.prototype_table_index]
                        if generated_table.prototype_table_index is not None
                        and generated_table.prototype_table_index < len(source_document.tables)
                        else matching_template_table(
                            list(source_document.tables),
                            columns=len(generated_table.headers),
                        )
                    ),
                    table,
                    profile=profile,
                )
            for generated_image in section.images:
                asset = image_assets.get(generated_image.asset_id)
                if asset is None:
                    raise TemplateInvalidError(
                        f"结构化图片缺少已确认文件：{generated_image.block_key}"
                    )
                _add_generated_image(
                    document,
                    asset=asset,
                    caption=generated_image.caption,
                    alt_text=generated_image.alt_text,
                    profile=profile,
                )

        page_summary = ensure_page_number_fields(document, profile=profile)
        if not page_summary.valid:
            raise TemplateInvalidError("页码字段校验失败，当前页或总页数字段不正确")
        output = BytesIO()
        document.save(output)
        content = output.getvalue()
        xml_text = _docx_xml_text(content)
        if "{{" in xml_text or "{%" in xml_text:
            raise TemplateInvalidError("渲染结果仍包含模板占位符")

        output_name = Path(request.template.filename).stem + "-generated.docx"
        return RenderedArtifact(
            filename=output_name,
            media_type=self._media_type,
            content=content,
            sha256=sha256(content).hexdigest(),
        )
