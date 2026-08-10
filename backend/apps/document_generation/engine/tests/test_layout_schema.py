from io import BytesIO

from docx import Document

from apps.document_generation.engine.layout_schema import infer_template_layout_schema


def test_layout_schema_assigns_semantic_section_without_recognized_outline() -> None:
    document = Document()
    document.add_paragraph("客户历史正文（不作为生成内容）")
    table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(("危险点", "控制措施", "责任人")):
        table.rows[0].cells[index].text = value
    output = BytesIO()
    document.save(output)

    schema = infer_template_layout_schema(output.getvalue())

    assert schema["table_slots"] == [
        {
            "block_key": "hazard_controls",
            "section_code": "risk_identification",
            "prototype_table_index": 0,
            "headers": ["危险点", "控制措施", "责任人"],
            "column_count": 3,
        }
    ]
    assert {slot["block_key"] for slot in schema["image_slots"]} == {
        "rescue_route",
        "height_escape_plan",
        "height_rescue_plan",
    }


def test_layout_schema_uses_approved_default_slots_for_tableless_customer_template() -> None:
    document = Document()
    for heading in (
        "第二章 组织措施",
        "第三章 安全措施",
        "第五章 技术措施",
        "第六章 施工方案",
        "第七章 现场处置方案",
    ):
        document.add_paragraph(heading)
    output = BytesIO()
    document.save(output)

    schema = infer_template_layout_schema(output.getvalue())

    assert schema["version"] == 2
    assert schema["uses_approved_default_table_style"] is True
    slots = {slot["block_key"]: slot for slot in schema["table_slots"]}
    assert slots["personnel_information"]["section_code"] == "organization_measures"
    assert slots["hazard_controls"]["section_code"] == "safety_measures"
    assert slots["work_process_points"]["section_code"] == "construction_plan"
    assert all(slot["style_source"] == "approved_default_v1" for slot in slots.values())
