from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from apps.document_generation.engine.contracts import (
    ParsedBlockType,
    SourceDocument,
)
from apps.document_generation.engine.errors import (
    SourcePurposeMismatchError,
    SourceUnsupportedError,
)
from apps.document_generation.engine.parsing import EntrySourceParser

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_source_bytes() -> bytes:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "重复页眉"
    document.add_paragraph("入场四措两案", style="Title")
    document.add_heading("一、工程概况", level=1)
    document.add_paragraph("")
    document.add_paragraph("计划检测12台风电机组。")
    document.add_paragraph("人员入场前完成安全交底。", style="List Bullet")
    document.add_paragraph("1")
    document.add_heading("1.1 工作范围", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "计划内容"
    table.cell(1, 0).text = "检测对象"
    table.cell(1, 1).text = "塔筒焊缝"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _result_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("现场技术资料", style="Title")
    document.add_paragraph("检测结论：存在需要处理的缺陷。")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text_pdf_bytes() -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    resources = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    page[NameObject("/Resources")] = resources
    stream = DecodedStreamObject()
    stream.set_data(
        b"BT /F1 12 Tf 72 720 Td (1. Project Overview) Tj "
        b"0 -20 Td (Planned entry inspection work) Tj ET"
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def test_docx_parser_preserves_headings_lists_tables_and_source_locations() -> None:
    parser = EntrySourceParser()
    content = _docx_source_bytes()
    source = SourceDocument(
        document_version_id=201,
        filename="入场四措两案.docx",
        mime_type=DOCX_MIME,
        content=content,
    )

    first = parser.parse(source)
    second = parser.parse(source)

    assert first.content_sha256 == second.content_sha256
    assert first.blocks == second.blocks
    assert "重复页眉" not in "\n".join(block.text for block in first.blocks)
    assert "1" not in [block.text for block in first.blocks]
    assert any(block.block_type == ParsedBlockType.LIST_ITEM for block in first.blocks)
    table = next(block for block in first.blocks if block.block_type == ParsedBlockType.TABLE)
    assert table.rows[0] == ("项目", "计划内容")
    assert table.rows[1] == ("检测对象", "塔筒焊缝")
    work_scope = next(block for block in first.blocks if block.text == "1.1 工作范围")
    assert work_scope.heading_path == ("一、工程概况", "1.1 工作范围")
    assert work_scope.locator.paragraph_index is not None


def test_pdf_parser_preserves_page_location() -> None:
    source = SourceDocument(
        document_version_id=202,
        filename="entry-requirements.pdf",
        mime_type="application/pdf",
        content=_text_pdf_bytes(),
    )

    parsed = EntrySourceParser().parse(source)

    assert parsed.title == "1. Project Overview"
    assert all(block.locator.page == 1 for block in parsed.blocks)
    assert any("Planned entry inspection work" in block.text for block in parsed.blocks)


def test_scanned_pdf_without_text_is_rejected() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)
    source = SourceDocument(
        document_version_id=203,
        filename="scanned.pdf",
        mime_type="application/pdf",
        content=output.getvalue(),
    )

    with pytest.raises(SourceUnsupportedError, match="扫描件"):
        EntrySourceParser().parse(source)


def test_legacy_doc_is_rejected() -> None:
    source = SourceDocument(
        document_version_id=204,
        filename="legacy.doc",
        mime_type="application/msword",
        content=b"legacy",
    )

    with pytest.raises(SourceUnsupportedError, match="暂不支持"):
        EntrySourceParser().parse(source)


@pytest.mark.parametrize("filename", ["检测报告.docx", "完工报告.pdf", "竣工资料.docx"])
def test_report_or_completion_source_is_rejected(filename: str) -> None:
    source = SourceDocument(
        document_version_id=205,
        filename=filename,
        mime_type=DOCX_MIME,
        content=_docx_source_bytes(),
    )

    with pytest.raises(SourcePurposeMismatchError):
        EntrySourceParser().parse(source)


def test_result_content_is_rejected_even_without_report_filename() -> None:
    source = SourceDocument(
        document_version_id=206,
        filename="现场技术资料.docx",
        mime_type=DOCX_MIME,
        content=_result_docx_bytes(),
    )

    with pytest.raises(SourcePurposeMismatchError):
        EntrySourceParser().parse(source)
