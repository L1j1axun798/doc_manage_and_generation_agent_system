from __future__ import annotations

import base64
import hashlib
from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from apps.document_generation.engine.contracts import (
    ConfirmedFact,
    GeneratedSection,
    GeneratedImage,
    GeneratedTable,
    RenderRequest,
    RenderImageAsset,
    SourceLocator,
    TemplateDocument,
)
from apps.document_generation.engine.docx_formatting import page_number_field_summary
from apps.document_generation.engine.errors import (
    SourcePurposeMismatchError,
    TemplateInvalidError,
)
from apps.document_generation.engine.rendering import (
    DocxTemplateRenderer,
    infer_template_section_order,
)


def _template_bytes() -> bytes:
    document = Document()
    document.add_heading("{{ project_name }}四措两案", level=0)
    document.add_paragraph("计划数量：{{ planned_inspection_quantity }}台")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _facts() -> tuple[ConfirmedFact, ...]:
    locator = SourceLocator(paragraph_index=1)
    return (
        ConfirmedFact(
            field="project_name",
            value="示例风电场入场项目",
            value_type="string",
            source_document_version_id=1,
            locator=locator,
            confidence=1,
            confirmed_by=1,
        ),
        ConfirmedFact(
            field="planned_inspection_quantity",
            value=12,
            value_type="integer",
            source_document_version_id=1,
            locator=locator,
            confidence=1,
            confirmed_by=1,
        ),
    )


def _append_field(paragraph, instruction: str, result: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_text.text = f" {instruction} \\* MERGEFORMAT "
    instruction_run._r.append(instruction_text)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(result)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def test_template_preflight_and_render_produce_openable_docx() -> None:
    template = TemplateDocument(
        template_id="tpl-1",
        filename="四措两案模板.docx",
        content=_template_bytes(),
        required_placeholders=("project_name", "planned_inspection_quantity"),
    )
    renderer = DocxTemplateRenderer()

    preflight = renderer.preflight(
        template,
        {
            "project_name": "示例风电场入场项目",
            "planned_inspection_quantity": 12,
        },
    )
    artifact = renderer.render(
        RenderRequest(
            template=template,
            facts=_facts(),
            sections=(
                GeneratedSection(
                    section_code="safety_measures",
                    title="安全措施",
                    paragraphs=("入场前完成安全交底。",),
                    lists=(("核验人员资质", "核验安全带"),),
                    tables=(
                        GeneratedTable(
                            headers=("风险", "措施"),
                            rows=(("高处作业", "使用双钩安全带"),),
                        ),
                    ),
                ),
            ),
        )
    )

    assert preflight.valid is True
    assert set(preflight.declared_placeholders) == {
        "project_name",
        "planned_inspection_quantity",
    }
    rendered = Document(BytesIO(artifact.content))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "示例风电场入场项目四措两案" in text
    assert "计划数量：12台" in text
    assert "安全措施" in text
    assert "入场前完成安全交底。" in text
    assert rendered.tables[-1].rows[0].cells[0].text == "风险"
    with ZipFile(BytesIO(artifact.content)) as archive:
        settings_xml = archive.read("word/settings.xml").decode()
        document_xml = archive.read("word/document.xml").decode()
    assert "w:updateFields" in settings_xml
    assert "{{" not in document_xml


def test_confirmed_image_renders_with_caption_alt_text_and_adequate_dpi() -> None:
    image = Image.new("RGB", (1600, 900), color=(235, 245, 255))
    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")
    content = image_buffer.getvalue()
    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=TemplateDocument(
                template_id="image-template",
                filename="entry-plan.docx",
                content=_template_bytes(),
            ),
            facts=_facts(),
            sections=(
                GeneratedSection(
                    section_code="emergency_plan",
                    title="应急预案",
                    images=(
                        GeneratedImage(
                            block_key="rescue_route",
                            asset_id="route-1",
                            title="救援路线图",
                            caption="风场至最近医院救援路线",
                            alt_text="风场至最近医院的驾车救援路线图",
                            insertion_reason="用户确认路线后冻结",
                            source_kind="rescue_route",
                        ),
                    ),
                ),
            ),
            image_assets=(
                RenderImageAsset(
                    asset_id="route-1",
                    filename="route.png",
                    media_type="image/png",
                    content=content,
                    sha256=hashlib.sha256(content).hexdigest(),
                    width_px=1600,
                    height_px=900,
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    assert "风场至最近医院救援路线" in [paragraph.text for paragraph in rendered.paragraphs]
    with ZipFile(BytesIO(artifact.content)) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "风场至最近医院的驾车救援路线图" in document_xml


def test_template_missing_required_placeholder_is_rejected() -> None:
    template = TemplateDocument(
        template_id="tpl-2",
        filename="四措两案模板.docx",
        content=_template_bytes(),
        required_placeholders=("project_name", "site_name"),
    )
    renderer = DocxTemplateRenderer()

    preflight = renderer.preflight(template, {"project_name": "项目"})

    assert preflight.valid is False
    assert set(preflight.missing_placeholders) == {
        "planned_inspection_quantity",
        "site_name",
    }
    with pytest.raises(TemplateInvalidError, match="模板预检失败"):
        renderer.render(
            RenderRequest(
                template=template,
                facts=_facts()[:1],
                sections=(),
            )
        )


def test_report_template_is_rejected() -> None:
    template = TemplateDocument(
        template_id="tpl-report",
        filename="检测报告模板.docx",
        content=_template_bytes(),
    )

    with pytest.raises(SourcePurposeMismatchError):
        DocxTemplateRenderer().preflight(template, {})


def test_filled_style_baseline_does_not_leak_historical_body_or_headers() -> None:
    historical = Document()
    cover = historical.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run("历史风电场项目四措两案")
    cover_run.font.size = Pt(22)
    body = historical.add_paragraph("历史甲方专有施工正文")
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Pt(30)
    body.paragraph_format.line_spacing = 1.75
    body.paragraph_format.space_after = Pt(6)
    body_run = body.runs[0]
    body_run.font.size = Pt(16)
    body_run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"),
        "仿宋_GB2312",
    )
    historical.add_picture(
        BytesIO(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMA"
                "ASsJTYQAAAAASUVORK5CYII="
            )
        )
    )
    historical.sections[0].header.paragraphs[0].text = "历史甲方页眉"
    historical.sections[0].footer.paragraphs[0].text = "历史项目页脚"
    appendix = historical.add_section(WD_SECTION.NEW_PAGE)
    appendix.orientation = WD_ORIENT.LANDSCAPE
    appendix.page_width, appendix.page_height = (
        appendix.page_height,
        appendix.page_width,
    )
    output = BytesIO()
    historical.save(output)
    template = TemplateDocument(
        template_id="tpl-style-only",
        filename="已批准样式基线.docx",
        content=output.getvalue(),
    )
    renderer = DocxTemplateRenderer()

    preflight = renderer.preflight(
        template,
        {
            "project_name": "当前风电场入场项目",
            "client_name": "当前委托方",
        },
    )
    current_project_fact = _facts()[0].model_copy(update={"value": "当前风电场入场项目"})
    client_fact = current_project_fact.model_copy(
        update={
            "field": "client_name",
            "value": "当前委托方",
        }
    )
    artifact = renderer.render(
        RenderRequest(
            template=template,
            facts=(current_project_fact, client_fact),
            sections=(
                GeneratedSection(
                    section_code="overview",
                    title="工程概况",
                    paragraphs=("当前项目计划开展入场检测。",),
                ),
            ),
        )
    )

    assert preflight.valid is True
    assert preflight.declared_placeholders == ()
    assert preflight.warnings
    rendered = Document(BytesIO(artifact.content))
    visible_text = "\n".join(
        [
            *(paragraph.text for paragraph in rendered.paragraphs),
            *(
                paragraph.text
                for section in rendered.sections
                for paragraph in (*section.header.paragraphs, *section.footer.paragraphs)
            ),
        ]
    )
    assert "当前风电场入场项目" in visible_text
    assert "当前委托方" in visible_text
    assert "当前项目计划开展入场检测" in visible_text
    assert "历史风电场" not in visible_text
    assert "历史甲方" not in visible_text
    assert "历史项目" not in visible_text
    assert rendered.sections[0].orientation == WD_ORIENT.PORTRAIT
    generated_body = next(
        paragraph
        for paragraph in rendered.paragraphs
        if "当前项目计划开展入场检测" in paragraph.text
    )
    assert generated_body.runs[0].font.size == Pt(16)
    assert (
        generated_body.runs[0]._element.get_or_add_rPr().get_or_add_rFonts().get(qn("w:eastAsia"))
        == "仿宋_GB2312"
    )
    assert generated_body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert generated_body.paragraph_format.first_line_indent == Pt(32)
    assert generated_body.paragraph_format.line_spacing == 1.5
    assert generated_body._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
    assert generated_body.paragraph_format.space_after == Pt(6)
    with ZipFile(BytesIO(artifact.content)) as archive:
        assert not any(name.startswith("word/media/") for name in archive.namelist())


def test_style_only_baseline_without_english_list_or_table_styles_still_renders() -> None:
    baseline = Document()
    for style_name in ("List Bullet", "Table Grid"):
        style = baseline.styles[style_name]
        style._element.getparent().remove(style._element)
    output = BytesIO()
    baseline.save(output)
    template = TemplateDocument(
        template_id="tpl-minimal-styles",
        filename="客户四措两案样式基线.docx",
        content=output.getvalue(),
    )

    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=template,
            facts=_facts()[:1],
            sections=(
                GeneratedSection(
                    section_code="safety_measures",
                    title="安全措施",
                    lists=(("核验人员资质",),),
                    tables=(
                        GeneratedTable(
                            headers=("风险", "措施"),
                            rows=(("高处作业", "使用双钩安全带"),),
                        ),
                    ),
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    assert "（1）核验人员资质" in [paragraph.text for paragraph in rendered.paragraphs]
    assert rendered.tables[-1].rows[1].cells[1].text == "使用双钩安全带"


def test_generated_content_uses_detected_template_format_and_removes_markdown() -> None:
    baseline = Document()
    baseline.add_heading("{{ project_name }}四措两案", level=0)
    chapter = baseline.add_heading("一、模板章节格式", level=2)
    chapter.runs[0].font.name = "黑体"
    chapter.runs[0].font.size = Pt(18)
    chapter.paragraph_format.space_before = Pt(12)
    chapter.paragraph_format.space_after = Pt(6)
    subheading = baseline.add_paragraph("1.1模板小标题")
    subheading.runs[0].font.name = "黑体"
    subheading.runs[0].font.size = Pt(14)
    subheading.runs[0].bold = True
    subheading.paragraph_format.space_before = Pt(6)
    body = baseline.add_paragraph("这是用于识别模板正文格式的完整叙述样例文本。")
    body.runs[0].font.name = "仿宋_GB2312"
    body.runs[0].font.size = Pt(14)
    body.paragraph_format.first_line_indent = Pt(10)
    body.paragraph_format.line_spacing = 1.25
    body.paragraph_format.space_after = Pt(5)
    leaked_outline_level = OxmlElement("w:outlineLvl")
    leaked_outline_level.set(qn("w:val"), "1")
    body._p.get_or_add_pPr().append(leaked_outline_level)
    second_body = baseline.add_paragraph("第二个正文样例用于确认局部加粗不会覆盖典型正文格式。")
    second_body.runs[0].font.name = "仿宋_GB2312"
    second_body.runs[0].font.size = Pt(14)
    second_body.paragraph_format.first_line_indent = Pt(10)
    second_body.paragraph_format.line_spacing = 1.25
    second_body.paragraph_format.space_after = Pt(5)
    outlier = baseline.add_paragraph("单个很长的历史强调段落" * 20)
    outlier.runs[0].font.name = "黑体"
    outlier.runs[0].font.size = Pt(18)
    outlier.runs[0].bold = True
    list_sample = baseline.add_paragraph("1）模板清单格式")
    list_sample.runs[0].font.name = "仿宋_GB2312"
    list_sample.runs[0].font.size = Pt(14)
    list_sample.paragraph_format.left_indent = Pt(14)
    source_table = baseline.add_table(rows=2, cols=2)
    source_table.columns[0].width = Inches(1.3)
    source_table.columns[1].width = Inches(4.7)
    source_table.cell(0, 0).text = "模板表头一"
    source_table.cell(0, 1).text = "模板表头二"
    source_table.cell(1, 0).text = "模板值一"
    source_table.cell(1, 1).text = "模板值二"
    source_table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(10.5)
    output = BytesIO()
    baseline.save(output)

    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=TemplateDocument(
                template_id="format-profile",
                filename="格式识别模板.docx",
                content=output.getvalue(),
                required_placeholders=("project_name",),
            ),
            facts=_facts()[:1],
            sections=(
                GeneratedSection(
                    section_code="safety_measures",
                    title="**安全措施**",
                    paragraphs=(
                        "### 1.1生成小标题",
                        "**生成正文**必须服从模板，不保留`模型标记`。",
                    ),
                    lists=(("- **核验人员资质**", "• 核验安全带"),),
                    tables=(
                        GeneratedTable(
                            headers=("**风险**", "措施"),
                            rows=(("`高处作业`", "使用双钩安全带"),),
                        ),
                    ),
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    generated_heading = next(p for p in rendered.paragraphs if p.text == "一、安全措施")
    generated_subheading = next(p for p in rendered.paragraphs if p.text == "1.1生成小标题")
    generated_body = next(p for p in rendered.paragraphs if p.text.startswith("生成正文"))
    generated_list = next(p for p in rendered.paragraphs if p.text == "1）核验人员资质")
    assert generated_heading.style.name == chapter.style.name
    assert generated_heading.runs[0].font.size == Pt(18)
    assert generated_heading.paragraph_format.space_after == Pt(6)
    assert generated_subheading.runs[0].font.size == Pt(14)
    assert generated_subheading.runs[0].bold is True
    assert generated_body.runs[0].font.size == Pt(14)
    assert generated_body.paragraph_format.first_line_indent == Pt(28)
    assert generated_body.paragraph_format.line_spacing == 1.5
    assert generated_body._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
    assert generated_body.paragraph_format.space_after == Pt(5)
    assert generated_body._p.pPr.find(qn("w:outlineLvl")) is None
    subheading_indentation = generated_subheading._p.pPr.ind
    assert (
        subheading_indentation is None
        or subheading_indentation.get(qn("w:firstLineChars")) is None
    )
    assert generated_list.paragraph_format.left_indent == Pt(14)
    assert [cell.text for cell in rendered.tables[-1].rows[0].cells] == ["风险", "措施"]
    assert rendered.tables[-1].cell(0, 0).paragraphs[0].runs[0].font.size == Pt(10.5)
    visible_text = "\n".join(
        [
            *(paragraph.text for paragraph in rendered.paragraphs),
            *(cell.text for table in rendered.tables for row in table.rows for cell in row.cells),
        ]
    )
    assert all(marker not in visible_text for marker in ("**", "###", "`", "•"))


def test_bad_total_page_field_and_stale_results_are_repaired() -> None:
    baseline = Document()
    baseline.add_paragraph("第一章、项目简介")
    baseline.add_paragraph("历史项目正文")
    footer = baseline.sections[0].footer.paragraphs[0]
    footer.add_run("第 ")
    _append_field(footer, "PAGE", "17")
    footer.add_run(" 页 共 30 页")
    output = BytesIO()
    baseline.save(output)

    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=TemplateDocument(
                template_id="bad-page-fields",
                filename="甲方四措两案.docx",
                content=output.getvalue(),
            ),
            facts=_facts()[:1],
            sections=(
                GeneratedSection(
                    section_code="overview",
                    title="项目简介",
                    paragraphs=("当前项目正文",),
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    summary = page_number_field_summary(rendered)
    assert summary.valid is True
    assert summary.page_fields >= 1
    assert summary.total_page_fields >= 1
    with ZipFile(BytesIO(artifact.content)) as archive:
        footer_xml = archive.read("word/footer1.xml").decode()
        settings_xml = archive.read("word/settings.xml").decode()
    assert "NUMPAGES" in footer_xml
    assert ">17<" not in footer_xml
    assert ">30<" not in footer_xml
    assert settings_xml.count("w:updateFields") == 1


def test_customer_outline_controls_section_order_without_leaking_historical_cover() -> None:
    baseline = Document()
    baseline.add_paragraph("甲方封面标题")
    baseline.add_paragraph("第一章、项目简介")
    baseline.add_paragraph("历史项目概况")
    baseline.add_paragraph("第二章、安全措施")
    baseline.add_paragraph("历史安全正文")
    baseline.add_paragraph("第三章、施工方案")
    baseline.add_paragraph("历史施工正文")
    for style_name in ("Heading 2", "Heading 3"):
        style = baseline.styles[style_name]
        style._element.getparent().remove(style._element)
    output = BytesIO()
    baseline.save(output)

    assert infer_template_section_order(output.getvalue()) == (
        "overview",
        "safety_measures",
        "construction_plan",
    )

    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=TemplateDocument(
                template_id="customer-outline",
                filename="甲方四措两案.docx",
                content=output.getvalue(),
            ),
            facts=_facts()[:1],
            sections=(
                GeneratedSection(
                    section_code="overview",
                    title="系统概况标题",
                    paragraphs=("当前项目概况",),
                ),
                GeneratedSection(
                    section_code="safety_measures",
                    title="系统安全标题",
                    paragraphs=("当前安全正文",),
                ),
                GeneratedSection(
                    section_code="construction_plan",
                    title="系统施工标题",
                    paragraphs=("当前施工正文",),
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    texts = [paragraph.text for paragraph in rendered.paragraphs if paragraph.text.strip()]
    assert texts == [
        "示例风电场入场项目",
        "入场四措两案",
        "一、系统概况标题",
        "当前项目概况",
        "二、系统安全标题",
        "当前安全正文",
        "三、系统施工标题",
        "当前施工正文",
    ]
    assert "甲方封面标题" not in texts
    assert "历史项目概况" not in texts
    first_chapter = next(p for p in rendered.paragraphs if p.text == "一、系统概况标题")
    assert first_chapter.style.name == baseline.paragraphs[1].style.name


def test_multiple_subheadings_do_not_become_the_body_format_sample() -> None:
    baseline = Document()
    baseline.add_paragraph("第一章、项目简介")
    first_subheading = baseline.add_paragraph("1.1 项目概况")
    first_subheading.runs[0].bold = True
    second_subheading = baseline.add_paragraph("1.2 编制依据")
    second_subheading.runs[0].bold = True
    baseline.add_paragraph("第二章、安全措施")
    body_sample = baseline.add_paragraph("这是模板中用于识别普通正文格式的叙述样例。")
    body_sample.runs[0].font.size = Pt(12)
    body_sample.paragraph_format.first_line_indent = Pt(24)
    output = BytesIO()
    baseline.save(output)

    artifact = DocxTemplateRenderer().render(
        RenderRequest(
            template=TemplateDocument(
                template_id="multiple-subheadings",
                filename="甲方四措两案.docx",
                content=output.getvalue(),
            ),
            facts=_facts()[:1],
            sections=(
                GeneratedSection(
                    section_code="overview",
                    title="项目简介",
                    paragraphs=("当前项目正文不应加粗。",),
                ),
                GeneratedSection(
                    section_code="safety_measures",
                    title="安全措施",
                    paragraphs=("当前安全正文。",),
                ),
            ),
        )
    )

    rendered = Document(BytesIO(artifact.content))
    generated = next(p for p in rendered.paragraphs if p.text == "当前项目正文不应加粗。")
    assert generated.runs[0].bold is not True
    assert generated.paragraph_format.first_line_indent == Pt(24)
    assert generated.paragraph_format.line_spacing == 1.5
    assert generated._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
