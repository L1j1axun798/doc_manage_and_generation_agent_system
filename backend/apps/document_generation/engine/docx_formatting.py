from __future__ import annotations

import re
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from typing import Any

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.font import CT_RPr
from docx.oxml.text.parfmt import CT_PPr
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

MARKDOWN_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s*")
MARKDOWN_BLOCK_RE = re.compile(r"^\s*(?:>|[-*+•])\s+")
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)]\([^)]+\)")
MARKDOWN_TAG_RE = re.compile(r"</?(?:strong|b|em|i|p|br)\s*/?>", re.IGNORECASE)
LIST_ITEM_PREFIX_RE = re.compile(
    r"^\s*(?:[-*+•]\s+|[（(]?\d+(?:\.\d+)*[）).、．]\s*|[一二三四五六七八九十]+、\s*)"
)
VISIBLE_LIST_PREFIX_RE = re.compile(
    r"^\s*(?P<number>\d+)(?P<suffix>[）)、.．])(?P<space>\s*)"
)
CHAPTER_HEADING_RE = re.compile(
    r"^(?:第\s*[一二三四五六七八九十百0-9]+\s*章|[一二三四五六七八九十]+、)"
)
SUBHEADING_RE = re.compile(
    r"^(?:[（(][一二三四五六七八九十0-9]+[）)]|[一二三四五六七八九十]+、|"
    r"\d+(?:\.\d+){1,4}[、.．]?|\d+[、.．](?!\s))"
)
STATIC_PAGE_TOTAL_RE = re.compile(
    r"(?:当前|第)?\s*(?P<current>\d+)\s*页\s*共\s*(?P<total>\d+)\s*页"
)
HYBRID_STATIC_TOTAL_RE = re.compile(r"共\s*\d+\s*页")


@dataclass(frozen=True)
class ParagraphFormatSample:
    paragraph_properties: CT_PPr | None = None
    run_properties: CT_RPr | None = None


@dataclass(frozen=True)
class TemplateFormatProfile:
    cover: ParagraphFormatSample
    chapter_heading: ParagraphFormatSample
    subheading: ParagraphFormatSample
    body: ParagraphFormatSample
    list_item: ParagraphFormatSample
    footer_page: ParagraphFormatSample
    list_marker: str | None = None


@dataclass(frozen=True)
class PageNumberFieldSummary:
    page_fields: int
    total_page_fields: int
    invalid_total_fields: int
    static_page_totals: int

    @property
    def valid(self) -> bool:
        return (
            self.page_fields > 0
            and self.invalid_total_fields == 0
            and self.static_page_totals == 0
        )


def clean_generated_text(value: str, *, list_item: bool = False) -> str:
    """Remove model-oriented markup while preserving the generated wording."""

    text = value.replace("\r", "\n").replace("```", " ")
    text = MARKDOWN_HEADING_RE.sub("", text)
    text = MARKDOWN_BLOCK_RE.sub("", text)
    text = MARKDOWN_LINK_RE.sub(r"\1", text)
    text = MARKDOWN_TAG_RE.sub(" ", text)
    text = text.replace("**", "").replace("__", "").replace("~~", "").replace("`", "")
    if list_item:
        text = LIST_ITEM_PREFIX_RE.sub("", text)
    return " ".join(text.split())


def is_subheading_text(value: str) -> bool:
    text = clean_generated_text(value)
    if len(text) > 80 or text.endswith(("。", "；", ";", "！", "？")):
        return False
    return bool(SUBHEADING_RE.match(text))


def _paragraph_sample(paragraph: Paragraph | None) -> ParagraphFormatSample:
    if paragraph is None:
        return ParagraphFormatSample()
    paragraph_properties = (
        deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None
    )
    run_properties: CT_RPr | None = None
    for run in paragraph.runs:
        if run.text.strip() and run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break
    return ParagraphFormatSample(
        paragraph_properties=paragraph_properties,
        run_properties=run_properties,
    )


def _sample_signature(sample: ParagraphFormatSample) -> str:
    paragraph_xml = (
        sample.paragraph_properties.xml if sample.paragraph_properties is not None else ""
    )
    run_xml = sample.run_properties.xml if sample.run_properties is not None else ""
    return f"{paragraph_xml}\n{run_xml}"


def _dominant_sample(paragraphs: list[Paragraph]) -> ParagraphFormatSample:
    if not paragraphs:
        return ParagraphFormatSample()
    weights: Counter[str] = Counter()
    samples: dict[str, ParagraphFormatSample] = {}
    for paragraph in paragraphs:
        sample = _paragraph_sample(paragraph)
        signature = _sample_signature(sample)
        weights[signature] += 1
        samples.setdefault(signature, sample)
    return samples[weights.most_common(1)[0][0]]


def _looks_like_chapter(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
    if style_name.startswith("toc"):
        return False
    return style_name.startswith("heading 1") or style_name.startswith("heading 2") or bool(
        CHAPTER_HEADING_RE.match(" ".join(paragraph.text.split()))
    )


def _footer_paragraphs(document: DocumentObject) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for root in _footer_roots(document):
        for element in root.iter(qn("w:p")):
            if any(
                _field_command(node.text or "") == "PAGE"
                for node in element.iter(qn("w:instrText"))
            ):
                paragraphs.append(Paragraph(element, document))
    return paragraphs


def infer_template_format_profile(
    document: DocumentObject,
    *,
    chapter_indexes: set[int] | None = None,
) -> TemplateFormatProfile:
    nonempty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    explicit_chapters = chapter_indexes or set()
    chapter_pairs = [
        (index, paragraph)
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() and (index in explicit_chapters or _looks_like_chapter(paragraph))
    ]
    chapter_positions = [index for index, _ in chapter_pairs]
    chapters = [paragraph for _, paragraph in chapter_pairs]
    body_start = min(chapter_positions) + 1 if chapter_positions else 0
    body_scope = [
        paragraph
        for paragraph in document.paragraphs[body_start:]
        if paragraph.text.strip()
        and not _looks_like_chapter(paragraph)
        and not is_subheading_text(paragraph.text)
        and not (paragraph.style and paragraph.style.name.lower().startswith(("toc", "title")))
        and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER
    ]
    narrative = [
        paragraph
        for paragraph in body_scope
        if len(paragraph.text.strip()) >= 20
        and VISIBLE_LIST_PREFIX_RE.match(paragraph.text) is None
    ]
    body = without_outline_level(_dominant_sample(narrative or body_scope or nonempty))

    subheadings = [
        paragraph
        for paragraph in document.paragraphs[body_start:]
        if paragraph.text.strip() and is_subheading_text(paragraph.text)
    ]
    chapter_heading = _dominant_sample(chapters) if chapters else body
    subheading = _dominant_sample(subheadings) if subheadings else body

    list_candidates: list[tuple[Paragraph, str]] = []
    for paragraph in document.paragraphs[body_start:]:
        if is_subheading_text(paragraph.text):
            continue
        match = VISIBLE_LIST_PREFIX_RE.match(paragraph.text)
        if match is None:
            continue
        marker = "{index}" + match.group("suffix") + match.group("space")
        list_candidates.append((paragraph, marker))
    marker_counts = Counter(marker for _, marker in list_candidates)
    list_marker = marker_counts.most_common(1)[0][0] if marker_counts else "（{index}）"
    list_paragraphs = [
        paragraph for paragraph, marker in list_candidates if marker == list_marker
    ]
    list_item = without_outline_level(
        _dominant_sample(list_paragraphs) if list_paragraphs else body
    )

    cover_scope = document.paragraphs[:body_start] if body_start else document.paragraphs
    cover_candidates = [
        paragraph
        for paragraph in cover_scope
        if paragraph.text.strip() and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    ]
    cover_paragraph = max(
        cover_candidates,
        key=lambda paragraph: max(
            (
                run.font.size.pt
                for run in paragraph.runs
                if run.text.strip() and run.font.size is not None
            ),
            default=0,
        ),
        default=None,
    )
    cover = _paragraph_sample(cover_paragraph) if cover_paragraph else chapter_heading
    footer_candidates = _footer_paragraphs(document)
    footer_page = _dominant_sample(footer_candidates) if footer_candidates else body
    return TemplateFormatProfile(
        cover=cover,
        chapter_heading=chapter_heading,
        subheading=subheading,
        body=body,
        list_item=list_item,
        footer_page=footer_page,
        list_marker=list_marker,
    )


def _apply_run_properties(run: Run, properties: CT_RPr | None) -> None:
    if properties is None:
        return
    if run._r.rPr is not None:
        run._r.remove(run._r.rPr)
    run._r.insert(0, deepcopy(properties))


def apply_paragraph_format(paragraph: Paragraph, sample: ParagraphFormatSample) -> None:
    if sample.paragraph_properties is not None:
        if paragraph._p.pPr is not None:
            paragraph._p.remove(paragraph._p.pPr)
        paragraph._p.insert(0, deepcopy(sample.paragraph_properties))
    for run in paragraph.runs:
        _apply_run_properties(run, sample.run_properties)


def apply_body_paragraph_layout(paragraph: Paragraph) -> None:
    """Enforce the approved narrative layout without changing template typography."""

    font_size = next(
        (
            run.font.size
            for run in paragraph.runs
            if run.text.strip() and run.font.size is not None
        ),
        None,
    )
    if font_size is None and paragraph.style is not None:
        font_size = paragraph.style.font.size
    font_size_pt = font_size.pt if font_size is not None else 12

    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(font_size_pt * 2)
    paragraph_format.line_spacing = 1.5

    indentation = paragraph._p.get_or_add_pPr().get_or_add_ind()
    for attribute in ("hanging", "hangingChars"):
        qualified_name = qn(f"w:{attribute}")
        if qualified_name in indentation.attrib:
            del indentation.attrib[qualified_name]
    indentation.set(qn("w:firstLineChars"), "200")


def without_outline_level(sample: ParagraphFormatSample) -> ParagraphFormatSample:
    paragraph_properties = (
        deepcopy(sample.paragraph_properties)
        if sample.paragraph_properties is not None
        else None
    )
    if paragraph_properties is not None:
        outline_level = paragraph_properties.find(qn("w:outlineLvl"))
        if outline_level is not None:
            paragraph_properties.remove(outline_level)
    return ParagraphFormatSample(
        paragraph_properties=paragraph_properties,
        run_properties=(
            deepcopy(sample.run_properties) if sample.run_properties is not None else None
        ),
    )


def format_list_item(
    value: str,
    *,
    index: int,
    profile: TemplateFormatProfile,
) -> str:
    text = clean_generated_text(value, list_item=True)
    if profile.list_marker is None:
        return text
    return f"{profile.list_marker.format(index=index)}{text}"


def matching_template_table(
    tables: list[Table] | tuple[Table, ...],
    *,
    columns: int,
) -> Table | None:
    return next(
        (table for table in tables if table.rows and len(table.rows[0].cells) == columns),
        None,
    )


def apply_table_format(
    source: Table | None,
    target: Table,
    *,
    profile: TemplateFormatProfile,
) -> None:
    target.autofit = False
    if source is not None:
        if source._tbl.tblPr is not None:
            if target._tbl.tblPr is not None:
                target._tbl.remove(target._tbl.tblPr)
            target._tbl.insert(0, deepcopy(source._tbl.tblPr))
        if source._tbl.tblGrid is not None:
            if target._tbl.tblGrid is not None:
                target._tbl.remove(target._tbl.tblGrid)
            insert_at = 1 if target._tbl.tblPr is not None else 0
            target._tbl.insert(insert_at, deepcopy(source._tbl.tblGrid))

    for row_index, row in enumerate(target.rows):
        row_properties = row._tr.get_or_add_trPr()
        if row_index == 0 and row_properties.find(qn("w:tblHeader")) is None:
            row_properties.append(OxmlElement("w:tblHeader"))
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        source_row = None
        if source is not None and source.rows:
            source_row = source.rows[min(row_index, len(source.rows) - 1)]
        for column_index, cell in enumerate(row.cells):
            source_cell = None
            if source_row is not None and column_index < len(source_row.cells):
                source_cell = source_row.cells[column_index]
                if source_cell._tc.tcPr is not None:
                    if cell._tc.tcPr is not None:
                        cell._tc.remove(cell._tc.tcPr)
                    cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                source_paragraph = None
                if source_cell is not None and source_cell.paragraphs:
                    source_paragraph = source_cell.paragraphs[
                        min(paragraph_index, len(source_cell.paragraphs) - 1)
                    ]
                apply_paragraph_format(
                    paragraph,
                    _paragraph_sample(source_paragraph) if source_paragraph else profile.body,
                )


def _footer_roots(document: DocumentObject) -> list[Any]:
    roots: list[Any] = []
    for part in document.part.package.parts:
        if not str(part.partname).startswith("/word/footer"):
            continue
        element = getattr(part, "element", None)
        if element is not None:
            roots.append(element)
    return roots


def _field_command(instruction: str) -> str:
    normalized = " ".join(instruction.split())
    return normalized.split(" ", 1)[0].upper() if normalized else ""


def _set_instruction(node: Any, command: str) -> None:
    node.text = f" {command} \\* MERGEFORMAT "


def _replace_paragraph_content_with_page_fields(
    paragraph_element: Any,
    sample: ParagraphFormatSample,
) -> None:
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    if sample.paragraph_properties is not None:
        existing = paragraph_element.find(qn("w:pPr"))
        if existing is not None:
            paragraph_element.remove(existing)
        paragraph_element.insert(0, deepcopy(sample.paragraph_properties))

    def append_text(value: str) -> None:
        run = OxmlElement("w:r")
        if sample.run_properties is not None:
            run.append(deepcopy(sample.run_properties))
        text = OxmlElement("w:t")
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text.text = value
        run.append(text)
        paragraph_element.append(run)

    def append_field(command: str) -> None:
        for field_type in ("begin",):
            run = OxmlElement("w:r")
            if sample.run_properties is not None:
                run.append(deepcopy(sample.run_properties))
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), field_type)
            run.append(field)
            paragraph_element.append(run)
        instruction_run = OxmlElement("w:r")
        if sample.run_properties is not None:
            instruction_run.append(deepcopy(sample.run_properties))
        instruction = OxmlElement("w:instrText")
        instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instruction.text = f" {command} \\* MERGEFORMAT "
        instruction_run.append(instruction)
        paragraph_element.append(instruction_run)
        for field_type in ("separate",):
            run = OxmlElement("w:r")
            if sample.run_properties is not None:
                run.append(deepcopy(sample.run_properties))
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), field_type)
            run.append(field)
            paragraph_element.append(run)
        append_text("1")
        end_run = OxmlElement("w:r")
        if sample.run_properties is not None:
            end_run.append(deepcopy(sample.run_properties))
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        paragraph_element.append(end_run)

    append_text("第 ")
    append_field("PAGE")
    append_text(" 页 共 ")
    append_field("NUMPAGES")
    append_text(" 页")


def _repair_total_page_instructions(paragraph_element: Any) -> int:
    visible_prefix = ""
    repairs = 0
    for node in paragraph_element.iter():
        if node.tag == qn("w:t"):
            visible_prefix += node.text or ""
        elif node.tag == qn("w:instrText"):
            if _field_command(node.text or "") != "PAGE":
                continue
            if re.search(r"(?:共|/|\bof)\s*$", visible_prefix, re.IGNORECASE):
                _set_instruction(node, "NUMPAGES")
                repairs += 1
    return repairs


def _reset_page_field_results(paragraph_element: Any) -> None:
    fields: list[dict[str, object]] = []
    for node in paragraph_element.iter():
        if node.tag == qn("w:fldChar"):
            field_type = node.get(qn("w:fldCharType"))
            if field_type == "begin":
                fields.append({"command": "", "in_result": False, "written": False})
            elif field_type == "separate" and fields:
                fields[-1]["in_result"] = True
            elif field_type == "end" and fields:
                fields.pop()
        elif node.tag == qn("w:instrText") and fields:
            fields[-1]["command"] = _field_command(node.text or "")
        elif node.tag == qn("w:t") and fields:
            field = fields[-1]
            if field["command"] not in {"PAGE", "NUMPAGES"} or not field["in_result"]:
                continue
            node.text = "1" if not field["written"] else ""
            field["written"] = True


def _set_update_fields(document: DocumentObject) -> None:
    settings = document.settings.element
    update_fields = list(settings.iterchildren(qn("w:updateFields")))
    if update_fields:
        update_fields[0].set(qn("w:val"), "true")
        for duplicate in update_fields[1:]:
            settings.remove(duplicate)
        return
    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")
    settings.append(element)


def ensure_page_number_fields(
    document: DocumentObject,
    *,
    profile: TemplateFormatProfile,
) -> PageNumberFieldSummary:
    roots = _footer_roots(document)
    for root in roots:
        for paragraph_element in list(root.iter(qn("w:p"))):
            nested_paragraphs = [
                child
                for child in paragraph_element.iter(qn("w:p"))
                if child is not paragraph_element
            ]
            instructions = list(paragraph_element.iter(qn("w:instrText")))
            visible = "".join(node.text or "" for node in paragraph_element.iter(qn("w:t")))
            commands = [_field_command(node.text or "") for node in instructions]
            has_hybrid_total = (
                "PAGE" in commands
                and "NUMPAGES" not in commands
                and HYBRID_STATIC_TOTAL_RE.search(visible) is not None
            )
            if not nested_paragraphs and (
                (not instructions and STATIC_PAGE_TOTAL_RE.search(visible))
                or has_hybrid_total
            ):
                _replace_paragraph_content_with_page_fields(
                    paragraph_element,
                    _paragraph_sample(Paragraph(paragraph_element, document)),
                )
            else:
                _repair_total_page_instructions(paragraph_element)
                _reset_page_field_results(paragraph_element)

    summary = page_number_field_summary(document)
    if summary.page_fields == 0:
        footer = document.sections[0].footer
        paragraph = footer.paragraphs[0]
        _replace_paragraph_content_with_page_fields(paragraph._p, profile.footer_page)
    _set_update_fields(document)
    return page_number_field_summary(document)


def page_number_field_summary(document: DocumentObject) -> PageNumberFieldSummary:
    page_fields = 0
    total_page_fields = 0
    invalid_total_fields = 0
    static_page_totals = 0
    for root in _footer_roots(document):
        for paragraph_element in root.iter(qn("w:p")):
            instructions = [
                _field_command(node.text or "")
                for node in paragraph_element.iter(qn("w:instrText"))
            ]
            page_fields += instructions.count("PAGE")
            total_page_fields += instructions.count("NUMPAGES")
            visible_prefix = ""
            for node in paragraph_element.iter():
                if node.tag == qn("w:t"):
                    visible_prefix += node.text or ""
                elif node.tag == qn("w:instrText"):
                    command = _field_command(node.text or "")
                    if re.search(r"(?:共|/|\bof)\s*$", visible_prefix, re.IGNORECASE):
                        if command != "NUMPAGES":
                            invalid_total_fields += 1
            if not instructions:
                visible = "".join(
                    node.text or "" for node in paragraph_element.iter(qn("w:t"))
                )
                if STATIC_PAGE_TOTAL_RE.search(visible):
                    static_page_totals += 1
            else:
                visible = "".join(
                    node.text or "" for node in paragraph_element.iter(qn("w:t"))
                )
                if (
                    "PAGE" in instructions
                    and "NUMPAGES" not in instructions
                    and HYBRID_STATIC_TOTAL_RE.search(visible)
                ):
                    static_page_totals += 1
    return PageNumberFieldSummary(
        page_fields=page_fields,
        total_page_fields=total_page_fields,
        invalid_total_fields=invalid_total_fields,
        static_page_totals=static_page_totals,
    )


def paragraph_sample(paragraph: Paragraph | None) -> ParagraphFormatSample:
    return _paragraph_sample(paragraph)


def apply_run_properties(run: Run, properties: CT_RPr | None) -> None:
    _apply_run_properties(run, properties)


def footer_roots(document: DocumentObject) -> list[Any]:
    return _footer_roots(document)
