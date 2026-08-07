from __future__ import annotations

import re
from collections.abc import Iterator
from hashlib import sha256
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from .contracts import (
    ParsedBlock,
    ParsedBlockType,
    ParsedDocument,
    SourceDocument,
    SourceLocator,
)
from .errors import AgentError, SourcePurposeMismatchError, SourceUnsupportedError

USER_PROMPT_MIME_TYPE = "application/x-wind-doc-agent-prompt"

REPORT_FILENAME_MARKERS = (
    "检测报告",
    "试验报告",
    "验收报告",
    "完工报告",
    "竣工资料",
)
RESULT_CONTENT_MARKERS = (
    "检测结果表明",
    "经检测发现",
    "检测结论",
    "缺陷清单",
    "处理结果",
)
STRONG_RESULT_CONTENT_MARKERS = RESULT_CONTENT_MARKERS[:-1]
PURE_PAGE_NUMBER_RE = re.compile(r"^(?:第\s*)?\d+\s*(?:页)?$")
CHINESE_HEADING_RE = re.compile(r"^[一二三四五六七八九十百]+[、.．]\s*")
NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)(?:[、.．]|\s+)")


def _ensure_entry_filename(filename: str) -> None:
    lowered = filename.lower()
    if any(marker.lower() in lowered for marker in REPORT_FILENAME_MARKERS):
        raise SourcePurposeMismatchError("报告或竣工资料不能作为入场四措两案来源")


def _ensure_entry_content(title: str, text: str) -> None:
    if any(marker in title for marker in REPORT_FILENAME_MARKERS):
        raise SourcePurposeMismatchError("报告类文档不能作为入场四措两案来源")
    if contains_result_content(text):
        raise SourcePurposeMismatchError("来源正文属于完工结果或检测结论资料")


def contains_result_content(text: str) -> bool:
    return any(marker in text for marker in STRONG_RESULT_CONTENT_MARKERS) or (
        sum(marker in text for marker in RESULT_CONTENT_MARKERS) >= 2
    )


def _paragraph_style_name(paragraph: Paragraph) -> str:
    style = paragraph.style
    return (style.name if style is not None else "").strip()


def _heading_level(paragraph: Paragraph, text: str) -> int | None:
    style_name = _paragraph_style_name(paragraph)
    style_match = re.search(r"(?:Heading|标题)\s*(\d+)", style_name, re.IGNORECASE)
    if style_match:
        return max(1, min(int(style_match.group(1)), 9))
    if CHINESE_HEADING_RE.match(text):
        return 1
    numbered = NUMBERED_HEADING_RE.match(text)
    if numbered:
        return min(numbered.group(1).count(".") + 1, 9)
    return None


def _pdf_heading_level(text: str) -> int | None:
    if CHINESE_HEADING_RE.match(text):
        return 1
    numbered = NUMBERED_HEADING_RE.match(text)
    if numbered:
        return min(numbered.group(1).count(".") + 1, 9)
    return None


def _updated_heading_path(path: list[str], level: int, heading: str) -> tuple[str, ...]:
    del path[level - 1 :]
    while len(path) < level - 1:
        path.append("")
    path.append(heading)
    return tuple(item for item in path if item)


def _iter_docx_blocks(document: DocumentObject) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    style_name = _paragraph_style_name(paragraph).lower()
    paragraph_properties = paragraph._p.pPr
    return "list" in style_name or (
        paragraph_properties is not None and paragraph_properties.numPr is not None
    )


class DocxSourceParser:
    def parse(self, source: SourceDocument) -> ParsedDocument:
        _ensure_entry_filename(source.filename)
        try:
            document = Document(BytesIO(source.content))
        except Exception as exc:
            raise AgentError("SOURCE_PARSE_FAILED", "DOCX文件无法解析") from exc

        blocks: list[ParsedBlock] = []
        heading_path: list[str] = []
        title = ""
        paragraph_index = 0
        table_index = 0
        for item in _iter_docx_blocks(document):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                current_index = paragraph_index
                paragraph_index += 1
                if not text or PURE_PAGE_NUMBER_RE.fullmatch(text):
                    continue
                style_name = _paragraph_style_name(item).lower()
                if style_name.startswith("toc") or style_name.startswith("目录"):
                    continue
                level = _heading_level(item, text)
                if level is not None:
                    current_path = _updated_heading_path(heading_path, level, text)
                    block_type = ParsedBlockType.HEADING
                else:
                    current_path = tuple(value for value in heading_path if value)
                    block_type = (
                        ParsedBlockType.LIST_ITEM
                        if _is_list_paragraph(item)
                        else ParsedBlockType.PARAGRAPH
                    )
                if not title:
                    title = text
                blocks.append(
                    ParsedBlock(
                        block_id=f"{source.document_version_id}:p:{current_index}",
                        block_type=block_type,
                        text=text,
                        heading_path=current_path,
                        locator=SourceLocator(
                            heading_path=current_path,
                            paragraph_index=current_index,
                            text_quote=text[:200],
                        ),
                    )
                )
                continue

            rows = tuple(
                tuple(cell.text.strip() for cell in row.cells)
                for row in item.rows
                if any(cell.text.strip() for cell in row.cells)
            )
            current_table_index = table_index
            table_index += 1
            if not rows:
                continue
            current_path = tuple(value for value in heading_path if value)
            text = "\n".join(" | ".join(row) for row in rows)
            blocks.append(
                ParsedBlock(
                    block_id=f"{source.document_version_id}:t:{current_table_index}",
                    block_type=ParsedBlockType.TABLE,
                    text=text,
                    heading_path=current_path,
                    locator=SourceLocator(
                        heading_path=current_path,
                        table_index=current_table_index,
                        text_quote=text[:200],
                    ),
                    rows=rows,
                )
            )

        if not blocks:
            raise AgentError("SOURCE_PARSE_FAILED", "DOCX正文为空")
        combined_text = "\n".join(block.text for block in blocks)
        _ensure_entry_content(title, combined_text)
        return ParsedDocument(
            document_version_id=source.document_version_id,
            filename=source.filename,
            mime_type=source.mime_type,
            content_sha256=sha256(source.content).hexdigest(),
            title=title,
            blocks=tuple(blocks),
        )


class PdfSourceParser:
    def parse(self, source: SourceDocument) -> ParsedDocument:
        _ensure_entry_filename(source.filename)
        try:
            reader = PdfReader(BytesIO(source.content))
        except Exception as exc:
            raise AgentError("SOURCE_PARSE_FAILED", "PDF文件无法解析") from exc
        if reader.is_encrypted:
            raise SourceUnsupportedError("暂不支持加密PDF")

        blocks: list[ParsedBlock] = []
        heading_path: list[str] = []
        title = ""
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                raise AgentError("SOURCE_PARSE_FAILED", "PDF文本提取失败") from exc
            for line_index, raw_line in enumerate(page_text.splitlines()):
                text = raw_line.strip()
                if not text or PURE_PAGE_NUMBER_RE.fullmatch(text):
                    continue
                level = _pdf_heading_level(text)
                if level is not None:
                    current_path = _updated_heading_path(heading_path, level, text)
                    block_type = ParsedBlockType.HEADING
                else:
                    current_path = tuple(value for value in heading_path if value)
                    block_type = ParsedBlockType.PARAGRAPH
                if not title:
                    title = text
                blocks.append(
                    ParsedBlock(
                        block_id=f"{source.document_version_id}:pdf:{page_number}:{line_index}",
                        block_type=block_type,
                        text=text,
                        heading_path=current_path,
                        locator=SourceLocator(
                            heading_path=current_path,
                            page=page_number,
                            paragraph_index=line_index,
                            text_quote=text[:200],
                        ),
                    )
                )

        if not blocks:
            raise SourceUnsupportedError("PDF没有可提取文本，可能是扫描件")
        combined_text = "\n".join(block.text for block in blocks)
        _ensure_entry_content(title, combined_text)
        return ParsedDocument(
            document_version_id=source.document_version_id,
            filename=source.filename,
            mime_type=source.mime_type,
            content_sha256=sha256(source.content).hexdigest(),
            title=title,
            blocks=tuple(blocks),
        )


class EntrySourceParser:
    _docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    _pdf_mimes = {"application/pdf", "application/x-pdf"}

    def __init__(self) -> None:
        self.docx_parser = DocxSourceParser()
        self.pdf_parser = PdfSourceParser()

    def parse(self, source: SourceDocument) -> ParsedDocument:
        if source.mime_type == USER_PROMPT_MIME_TYPE:
            try:
                text = source.content.decode("utf-8").strip()
            except UnicodeDecodeError as exc:
                raise AgentError("SOURCE_PARSE_FAILED", "用户Prompt无法解析") from exc
            if not text:
                raise AgentError("SOURCE_PARSE_FAILED", "用户Prompt为空")
            prompt_lines = tuple(line.strip() for line in text.splitlines() if line.strip())
            return ParsedDocument(
                document_version_id=source.document_version_id,
                filename=source.filename,
                mime_type=source.mime_type,
                content_sha256=sha256(source.content).hexdigest(),
                title="用户本次编制要求",
                blocks=tuple(
                    ParsedBlock(
                        block_id=f"prompt:p:{index}",
                        block_type=ParsedBlockType.PARAGRAPH,
                        text=line,
                        locator=SourceLocator(
                            paragraph_index=index,
                            text_quote=line[:200],
                        ),
                    )
                    for index, line in enumerate(prompt_lines)
                ),
            )
        suffix = Path(source.filename).suffix.lower()
        if suffix == ".docx" or source.mime_type == self._docx_mime:
            return self.docx_parser.parse(source)
        if suffix == ".pdf" or source.mime_type.lower() in self._pdf_mimes:
            return self.pdf_parser.parse(source)
        raise SourceUnsupportedError(f"暂不支持来源格式：{suffix or source.mime_type}")
