from __future__ import annotations

from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from docx import Document as DocxDocument

from apps.audit.models import AuditLog
from apps.documents.models import Document, DocumentVersion
from apps.folders.models import Folder, PersonnelProfile
from apps.projects.models import Project, ProjectMember
from common.storage import LocalDocumentStorage

from ..jobs import run_generation_task
from ..models import (
    ApprovalStatus,
    DocumentTemplate,
    GeneratedSection,
    GenerationReview,
    GenerationTask,
    GenerationTraceEvent,
)

User = get_user_model()
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_user(username: str, role: str):
    return User.objects.create_user(
        username=username,
        password="Password123!",
        real_name=username,
        role=role,
    )


def docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for value in paragraphs:
        document.add_paragraph(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def docx_with_method_table_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("项目名称：风场入场项目")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "对变桨轴承采用相控阵超声无损探伤。"
    table.cell(1, 0).text = "对风电机组高强度螺栓采用相控阵超声探伤。"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def create_stored_version(
    *,
    root,
    actor,
    project,
    folder,
    title: str,
    filename: str,
    content: bytes,
    source_type: str = Document.SourceType.PROJECT_UPLOAD,
) -> DocumentVersion:
    storage = LocalDocumentStorage(root=root)
    stored = storage.save_uploaded_file(
        SimpleUploadedFile(filename, content, content_type=DOCX_MIME)
    )
    document = Document.objects.create(
        project=project,
        folder=folder,
        title=title,
        source_type=source_type,
        created_by=actor,
    )
    version = DocumentVersion.objects.create(
        document=document,
        version_number=1,
        original_filename=filename,
        content_type=DOCX_MIME,
        file_size=stored.size,
        sha256=stored.sha256,
        storage_path=stored.relative_path,
        uploaded_by=actor,
    )
    document.current_version = version
    document.save(update_fields=["current_version", "updated_at"])
    return version


def setup_generation_case(tmp_path):
    admin = make_user("admin", User.Role.SYSTEM_ADMIN)
    manager = make_user("manager", User.Role.PROJECT_MANAGER)
    viewer = make_user("viewer", User.Role.DATA_OPERATOR)
    project = Project.objects.create(name="风场入场项目", code="ENTRY-001", created_by=admin)
    ProjectMember.objects.create(
        project=project,
        user=manager,
        role=ProjectMember.Role.MANAGER,
        can_upload=True,
        can_download_restricted=True,
    )
    ProjectMember.objects.create(
        project=project,
        user=viewer,
        role=ProjectMember.Role.VIEWER,
    )
    staff_root = Folder.objects.create(
        name="人员资质",
        code="PUBLIC-STAFF",
        is_system_root=True,
        created_by=admin,
    )
    personnel_folder = Folder.objects.create(
        parent=staff_root,
        name="张三",
        created_by=admin,
    )
    PersonnelProfile.objects.create(
        folder=personnel_folder,
        gender=PersonnelProfile.Gender.MALE,
        id_card_number="110101199001010011",
        phone="13800000000",
        updated_by=admin,
    )
    technical_folder = Folder.objects.create(
        project=project,
        name="技术方案",
        code="PUBLIC-TECH-SOLUTION",
        created_by=admin,
    )
    entry_folder = Folder.objects.create(
        project=project,
        name="入场前置资料",
        code="PUBLIC-COMPLETION",
        created_by=admin,
    )
    report_folder = Folder.objects.create(
        project=project,
        name="报告模板",
        code="PUBLIC-REPORT-TEMPLATE",
        created_by=admin,
    )
    template_version = create_stored_version(
        root=tmp_path,
        actor=admin,
        project=project,
        folder=technical_folder,
        title="四措两案样式基线",
        filename="template.docx",
        content=docx_bytes("四措两案样式基线"),
    )
    source_version = create_stored_version(
        root=tmp_path,
        actor=manager,
        project=project,
        folder=entry_folder,
        title="入场任务通知",
        filename="entry-notice.docx",
        content=docx_bytes("项目名称：风场入场项目", "计划检测数量：12"),
        source_type=Document.SourceType.ENTRANCE_MATERIAL,
    )
    report_version = create_stored_version(
        root=tmp_path,
        actor=manager,
        project=project,
        folder=report_folder,
        title="检测报告模板",
        filename="report-template.docx",
        content=docx_bytes("报告模板"),
    )
    template = DocumentTemplate.objects.create(
        code="T001",
        version="v1",
        document_version=template_version,
        section_order=["overview"],
        required_fact_fields=["project_name"],
        is_active=True,
        approval_status=ApprovalStatus.APPROVED,
        approved_by=admin,
        created_by=admin,
    )
    return {
        "admin": admin,
        "manager": manager,
        "viewer": viewer,
        "project": project,
        "template": template,
        "source_version": source_version,
        "report_version": report_version,
        "technical_folder": technical_folder,
        "entry_folder": entry_folder,
        "personnel_folder": personnel_folder,
    }


def create_task_via_api(client, case):
    response = client.post(
        "/api/v1/document-generation/tasks/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "create-001",
            "facts": [
                {
                    "field": "project_name",
                    "value": "风场入场项目",
                    "value_type": "string",
                }
            ],
        },
        content_type="application/json",
    )
    assert response.status_code == 201
    return GenerationTask.objects.get(pk=response.json()["id"])


def prepare_confirmation_task(client, case):
    task = create_task_via_api(client, case)
    response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [case["source_version"].pk]},
        content_type="application/json",
    )
    assert response.status_code == 200
    task.status = GenerationTask.Status.NEEDS_CONFIRMATION
    task.progress = 20
    task.save(update_fields=["status", "progress", "updated_at"])
    return task


@pytest.mark.django_db(transaction=True)
def test_running_conversation_can_be_stopped_without_worker_overwrite(
    client,
    tmp_path,
    monkeypatch,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    task.status = GenerationTask.Status.GENERATING
    task.operation = GenerationTask.Operation.GENERATE
    task.progress = 68
    task.save(update_fields=["status", "operation", "progress", "updated_at"])
    stop_signals: list[str] = []
    monkeypatch.setattr(
        "apps.document_generation.services.stop_generation_job",
        lambda task_id: stop_signals.append(task_id) or "stop_requested",
    )

    response = client.post(f"/api/v1/document-generation/tasks/{task.pk}/stop/")

    task.refresh_from_db()
    assert response.status_code == 200
    assert response.json()["status"] == "cancelled"
    assert task.status == GenerationTask.Status.CANCELLED
    assert task.progress == 68
    assert task.completed_at is not None
    assert stop_signals == [str(task.pk)]
    assert task.reviews.filter(action=GenerationReview.Action.STOPPED).exists()
    assert task.workflow_events.filter(
        tool="cancel_generation_task",
        stage="cancelled",
        status="succeeded",
    ).exists()
    assert run_generation_task(str(task.pk)) == "skipped"
    task.refresh_from_db()
    assert task.status == GenerationTask.Status.CANCELLED


@pytest.mark.django_db(transaction=True)
def test_terminal_conversation_delete_is_soft_and_removes_private_draft(
    client,
    tmp_path,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    storage = LocalDocumentStorage(root=tmp_path)
    stored = storage.save_uploaded_file(
        SimpleUploadedFile("draft.docx", b"private draft", content_type=DOCX_MIME)
    )
    task.draft_storage_path = stored.relative_path
    task.save(update_fields=["draft_storage_path", "updated_at"])

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.delete(f"/api/v1/document-generation/tasks/{task.pk}/")

    task.refresh_from_db()
    assert response.status_code == 204
    assert task.deleted_at is not None
    assert task.deleted_by == case["manager"]
    assert not storage.exists(stored.relative_path)
    assert (
        client.get(
            "/api/v1/document-generation/tasks/",
            {"project": case["project"].pk},
        ).json()["count"]
        == 0
    )
    assert client.get(f"/api/v1/document-generation/tasks/{task.pk}/").status_code == 404
    assert AuditLog.objects.filter(
        action="document_generation.task.delete",
        resource_id=str(task.pk),
        result=AuditLog.Result.SUCCESS,
    ).exists()


@pytest.mark.django_db(transaction=True)
def test_running_conversation_must_be_stopped_before_delete(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    task.status = GenerationTask.Status.EXTRACTING
    task.save(update_fields=["status", "updated_at"])

    response = client.delete(f"/api/v1/document-generation/tasks/{task.pk}/")

    task.refresh_from_db()
    assert response.status_code == 409
    assert response.json()["code"] == "TASK_STILL_RUNNING"
    assert task.deleted_at is None


@pytest.mark.django_db(transaction=True)
def test_viewer_cannot_stop_or_delete_project_conversation(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    task.status = GenerationTask.Status.GENERATING
    task.operation = GenerationTask.Operation.GENERATE
    task.save(update_fields=["status", "operation", "updated_at"])
    client.force_login(case["viewer"])

    stop_response = client.post(f"/api/v1/document-generation/tasks/{task.pk}/stop/")
    task.status = GenerationTask.Status.DRAFT
    task.save(update_fields=["status", "updated_at"])
    delete_response = client.delete(f"/api/v1/document-generation/tasks/{task.pk}/")

    task.refresh_from_db()
    assert stop_response.status_code == 403
    assert delete_response.status_code == 403
    assert task.deleted_at is None


@pytest.mark.django_db(transaction=True)
def test_fact_confirmation_canonicalizes_shortened_quote_at_structural_position(
    client,
    tmp_path,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = prepare_confirmation_task(client, case)

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.put(
            f"/api/v1/document-generation/tasks/{task.pk}/facts/confirm/",
            {
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                        "source_document_version_id": case["source_version"].pk,
                        "locator": {
                            "paragraph_index": 0,
                            "text_quote": "项目名称：风场...",
                        },
                        "confidence": 1,
                    }
                ]
            },
            content_type="application/json",
        )

    task.refresh_from_db()
    assert response.status_code == 200
    assert task.status == GenerationTask.Status.READY
    assert task.facts_snapshot[0]["locator"]["text_quote"] == "项目名称：风场入场项目"


@pytest.mark.django_db(transaction=True)
def test_fact_confirmation_repairs_conflicting_method_locator_from_source_text(
    client,
    tmp_path,
):
    case = setup_generation_case(tmp_path)
    method_version = create_stored_version(
        root=tmp_path,
        actor=case["manager"],
        project=case["project"],
        folder=case["entry_folder"],
        title="相控阵检测方案",
        filename="method-plan.docx",
        content=docx_with_method_table_bytes(),
        source_type=Document.SourceType.ENTRANCE_MATERIAL,
    )
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    source_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [method_version.pk]},
        content_type="application/json",
    )
    assert source_response.status_code == 200
    task.status = GenerationTask.Status.NEEDS_CONFIRMATION
    task.progress = 20
    task.save(update_fields=["status", "progress", "updated_at"])

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.put(
            f"/api/v1/document-generation/tasks/{task.pk}/facts/confirm/",
            {
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                        "source_document_version_id": method_version.pk,
                        "locator": {
                            "paragraph_index": 0,
                            "text_quote": "项目名称：风场入场项目",
                        },
                        "confidence": 1,
                    },
                    {
                        "field": "inspection_method_codes",
                        "value": ["PAUT"],
                        "value_type": "list[string]",
                        "source_document_version_id": method_version.pk,
                        "locator": {
                            "paragraph_index": 110,
                            "table_index": 0,
                            "text_quote": (
                                "对变桨轴承采用相控阵超声无损探伤。\n"
                                "对风电机组高强度螺栓采用相控阵超声探伤。"
                            ),
                        },
                        "confidence": 1,
                    },
                ]
            },
            content_type="application/json",
        )

    task.refresh_from_db()
    assert response.status_code == 200
    assert task.status == GenerationTask.Status.READY
    method_fact = next(
        fact for fact in task.facts_snapshot if fact["field"] == "inspection_method_codes"
    )
    assert method_fact["locator"]["paragraph_index"] is None
    assert method_fact["locator"]["table_index"] == 0
    assert "相控阵超声" in method_fact["locator"]["text_quote"]


@pytest.mark.django_db(transaction=True)
def test_fact_confirmation_does_not_repair_unsupported_method_code(
    client,
    tmp_path,
):
    case = setup_generation_case(tmp_path)
    method_version = create_stored_version(
        root=tmp_path,
        actor=case["manager"],
        project=case["project"],
        folder=case["entry_folder"],
        title="相控阵检测方案",
        filename="method-plan.docx",
        content=docx_with_method_table_bytes(),
        source_type=Document.SourceType.ENTRANCE_MATERIAL,
    )
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    source_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [method_version.pk]},
        content_type="application/json",
    )
    assert source_response.status_code == 200
    task.status = GenerationTask.Status.NEEDS_CONFIRMATION
    task.progress = 20
    task.save(update_fields=["status", "progress", "updated_at"])

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.put(
            f"/api/v1/document-generation/tasks/{task.pk}/facts/confirm/",
            {
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                        "source_document_version_id": method_version.pk,
                        "locator": {
                            "paragraph_index": 0,
                            "text_quote": "项目名称：风场入场项目",
                        },
                        "confidence": 1,
                    },
                    {
                        "field": "inspection_method_codes",
                        "value": ["MT"],
                        "value_type": "list[string]",
                        "source_document_version_id": method_version.pk,
                        "locator": {
                            "paragraph_index": 110,
                            "table_index": 0,
                            "text_quote": "对变桨轴承采用相控阵超声无损探伤。",
                        },
                        "confidence": 1,
                    },
                ]
            },
            content_type="application/json",
        )

    task.refresh_from_db()
    assert response.status_code == 400
    assert response.json()["code"] == "FACT_EVIDENCE_INVALID"
    assert "系统已尝试自动修复来源定位" in response.json()["message"]
    assert task.status == GenerationTask.Status.NEEDS_CONFIRMATION


@pytest.mark.django_db(transaction=True)
def test_fact_confirmation_rejects_invalid_evidence_before_queueing_generation(
    client,
    tmp_path,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = prepare_confirmation_task(client, case)

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.put(
            f"/api/v1/document-generation/tasks/{task.pk}/facts/confirm/",
            {
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                        "source_document_version_id": case["source_version"].pk,
                        "locator": {
                            "paragraph_index": 99,
                            "text_quote": "不存在的来源",
                        },
                        "confidence": 1,
                    }
                ]
            },
            content_type="application/json",
        )

    task.refresh_from_db()
    assert response.status_code == 400
    assert response.json()["code"] == "FACT_EVIDENCE_INVALID"
    assert task.status == GenerationTask.Status.NEEDS_CONFIRMATION


@pytest.mark.django_db(transaction=True)
def test_pipeline_endpoint_atomically_creates_sources_and_queues_extraction(
    client,
    tmp_path,
    monkeypatch,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    queued_ids: list[str] = []
    monkeypatch.setattr(
        "apps.document_generation.queues.queue_generation_task",
        lambda task_id: queued_ids.append(task_id),
    )

    response = client.post(
        "/api/v1/document-generation/tasks/pipeline/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "pipeline-001",
            "document_version_ids": [case["source_version"].pk],
            "conversation_context": {
                "initial_message": "请重点核对入场人员分工后开始编制",
                "selected_personnel_ids": [case["personnel_folder"].pk],
            },
            "facts": [
                {
                    "field": "project_name",
                    "value": "风场入场项目",
                    "value_type": "string",
                }
            ],
        },
        content_type="application/json",
    )

    task = GenerationTask.objects.get(pk=response.json()["id"])
    assert response.status_code == 202
    assert task.status == GenerationTask.Status.EXTRACTING
    assert list(task.sources.values_list("document_version_id", flat=True)) == [
        case["source_version"].pk
    ]
    assert queued_ids == [str(task.pk)]
    assert task.conversation_context["initial_message"] == ("请重点核对入场人员分工后开始编制")
    assert task.conversation_context["personnel"] == [
        {
            "id": str(case["personnel_folder"].pk),
            "name": "张三",
            "gender": "male",
            "id_card_number": "110101199001010011",
            "phone": "13800000000",
            "job_title": "",
            "department": "",
            "contact": "13800000000",
            "certifications": [],
            "certificate_valid_until": None,
            "additional_info": {
                "personnel_folder_id": case["personnel_folder"].pk,
                "profile_complete": True,
            },
        }
    ]
    assert task.conversation_context["template"]["format_locked"] is True
    assert response.json()["conversation_context"] == task.conversation_context
    assert GenerationTraceEvent.objects.filter(
        task=task,
        tool="queue_fact_extraction",
    ).exists()


@pytest.mark.django_db
@pytest.mark.parametrize("actor_key", ["manager", "admin"])
def test_manager_and_admin_can_upload_self_service_template(
    client,
    tmp_path,
    actor_key,
):
    case = setup_generation_case(tmp_path)
    client.force_login(case[actor_key])

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.post(
            "/api/v1/document-generation/templates/",
            {
                "project_id": case["project"].pk,
                "file": SimpleUploadedFile(
                    f"client-{actor_key}.docx",
                    docx_bytes("甲方四措两案模板"),
                    content_type=DOCX_MIME,
                ),
            },
        )

    assert response.status_code == 201
    assert response.json()["sync_status"] == "synced"
    template = DocumentTemplate.objects.get(pk=response.json()["id"])
    assert template.is_active is True
    assert template.approval_status == ApprovalStatus.DRAFT
    assert template.field_mapping["self_service"] is True
    assert template.field_mapping["project_id"] == case["project"].pk
    assert template.document_version.document.folder == case["entry_folder"]

    listed = client.get(
        "/api/v1/document-generation/templates/",
        {"project_id": case["project"].pk},
    )
    assert listed.status_code == 200
    assert template.pk in {item["id"] for item in listed.json()}


@pytest.mark.django_db(transaction=True)
def test_self_service_template_can_be_used_without_approval(client, tmp_path, monkeypatch):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    monkeypatch.setattr(
        "apps.document_generation.queues.queue_generation_task",
        lambda task_id: None,
    )
    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        uploaded = client.post(
            "/api/v1/document-generation/templates/",
            {
                "project_id": case["project"].pk,
                "file": SimpleUploadedFile(
                    "client-direct.docx",
                    docx_bytes("甲方四措两案模板"),
                    content_type=DOCX_MIME,
                ),
            },
        )

        response = client.post(
            "/api/v1/document-generation/tasks/pipeline/",
            {
                "project_id": case["project"].pk,
                "template_id": uploaded.json()["id"],
                "idempotency_key": "self-service-template-pipeline",
                "document_version_ids": [case["source_version"].pk],
                "conversation_context": {"initial_message": "按甲方模板开始编制"},
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                    }
                ],
            },
            content_type="application/json",
        )

    assert uploaded.status_code == 201
    assert response.status_code == 202
    assert response.json()["template_id"] == uploaded.json()["id"]


@pytest.mark.django_db
def test_template_upload_skips_sync_when_entry_folder_is_missing(client, tmp_path):
    case = setup_generation_case(tmp_path)
    case["entry_folder"].is_active = False
    case["entry_folder"].save(update_fields=["is_active", "updated_at"])
    client.force_login(case["manager"])

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.post(
            "/api/v1/document-generation/templates/",
            {
                "project_id": case["project"].pk,
                "file": SimpleUploadedFile(
                    "client-no-entry.docx",
                    docx_bytes("甲方四措两案模板"),
                    content_type=DOCX_MIME,
                ),
            },
        )

    assert response.status_code == 201
    assert response.json()["sync_status"] == "folder_missing"
    template = DocumentTemplate.objects.get(pk=response.json()["id"])
    assert template.document_version.document.folder == case["technical_folder"]


@pytest.mark.django_db
def test_selecting_template_syncs_once_to_entry_folder(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    url = f'/api/v1/document-generation/templates/{case["template"].pk}/select/'

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        first = client.post(
            url,
            {"project_id": case["project"].pk},
            content_type="application/json",
        )
        second = client.post(
            url,
            {"project_id": case["project"].pk},
            content_type="application/json",
        )

    assert first.status_code == 200
    assert first.json()["sync_status"] == "synced"
    assert second.status_code == 200
    assert second.json()["sync_status"] == "already_present"
    assert Document.objects.filter(
        folder=case["entry_folder"],
        description__contains="四措两案 Agent 甲方模板",
    ).count() == 1
    synchronized_document = Document.objects.get(
        folder=case["entry_folder"],
        description__contains="四措两案 Agent 甲方模板",
    )
    task = create_task_via_api(client, case)
    source_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [synchronized_document.current_version_id]},
        content_type="application/json",
    )
    assert source_response.status_code == 400
    assert source_response.json()["code"] == "SOURCE_PURPOSE_MISMATCH"


@pytest.mark.django_db
def test_available_personnel_comes_from_public_staff_folders(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])

    response = client.get(
        "/api/v1/document-generation/tasks/available-personnel/",
        {"project_id": case["project"].pk},
    )

    assert response.status_code == 200
    assert response.json() == [
        {
            "id": str(case["personnel_folder"].pk),
            "folder_id": case["personnel_folder"].pk,
            "name": "张三",
            "gender": "male",
            "gender_display": "男",
            "id_card_number": "110101199001010011",
            "phone": "13800000000",
            "profile_complete": True,
            "updated_at": response.json()[0]["updated_at"],
        }
    ]


@pytest.mark.django_db(transaction=True)
def test_pipeline_rejects_personnel_outside_public_staff_list(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])

    response = client.post(
        "/api/v1/document-generation/tasks/pipeline/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "pipeline-invalid-personnel",
            "document_version_ids": [case["source_version"].pk],
            "conversation_context": {
                "initial_message": "开始编制",
                "selected_personnel_ids": [999999],
            },
        },
        content_type="application/json",
    )

    assert response.status_code == 400
    assert response.json()["code"] == "PERSONNEL_INVALID"
    assert (
        GenerationTask.objects.filter(idempotency_key="pipeline-invalid-personnel").exists()
        is False
    )


@pytest.mark.django_db
def test_create_is_idempotent_and_project_isolation_is_enforced(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)

    same = client.post(
        "/api/v1/document-generation/tasks/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "create-001",
            "facts": [
                {
                    "field": "project_name",
                    "value": "风场入场项目",
                    "value_type": "string",
                }
            ],
        },
        content_type="application/json",
    )
    conflict = client.post(
        "/api/v1/document-generation/tasks/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "create-001",
            "facts": [
                {
                    "field": "project_name",
                    "value": "另一个项目",
                    "value_type": "string",
                }
            ],
        },
        content_type="application/json",
    )
    client.force_login(case["viewer"])
    hidden = client.get(f"/api/v1/document-generation/tasks/{task.pk}/")

    assert same.status_code == 200
    assert same.json()["id"] == str(task.pk)
    assert conflict.status_code == 409
    assert conflict.json()["code"] == "IDEMPOTENCY_CONFLICT"
    assert hidden.status_code == 200
    assert AuditLog.objects.filter(action="document_generation.task.create").count() == 1


@pytest.mark.django_db
def test_viewer_cannot_create_and_report_source_is_rejected(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["viewer"])
    denied = client.post(
        "/api/v1/document-generation/tasks/",
        {
            "project_id": case["project"].pk,
            "template_id": case["template"].pk,
            "idempotency_key": "viewer-create",
        },
        content_type="application/json",
    )
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    blocked = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [case["report_version"].pk]},
        content_type="application/json",
    )

    assert denied.status_code == 404
    assert blocked.status_code == 400
    assert blocked.json()["code"] == "SOURCE_PURPOSE_MISMATCH"


@pytest.mark.django_db
def test_ordinary_and_other_project_documents_cannot_be_added_as_sources(client, tmp_path):
    case = setup_generation_case(tmp_path)
    ordinary_version = create_stored_version(
        root=tmp_path,
        actor=case["admin"],
        project=case["project"],
        folder=case["technical_folder"],
        title="普通项目通知",
        filename="ordinary-notice.docx",
        content=docx_bytes("普通项目资料"),
    )
    other_project = Project.objects.create(
        name="其他项目",
        code="ENTRY-002",
        created_by=case["admin"],
    )
    other_entry_folder = Folder.objects.create(
        project=other_project,
        name="入场前置资料",
        code="PUBLIC-COMPLETION",
        created_by=case["admin"],
    )
    other_entry_version = create_stored_version(
        root=tmp_path,
        actor=case["admin"],
        project=other_project,
        folder=other_entry_folder,
        title="其他项目入场资料",
        filename="other-entry.docx",
        content=docx_bytes("其他项目"),
        source_type=Document.SourceType.ENTRANCE_MATERIAL,
    )
    client.force_login(case["admin"])
    task = create_task_via_api(client, case)

    ordinary_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [ordinary_version.pk]},
        content_type="application/json",
    )
    other_project_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [other_entry_version.pk]},
        content_type="application/json",
    )

    assert ordinary_response.status_code == 400
    assert ordinary_response.json()["code"] == "SOURCE_PURPOSE_MISMATCH"
    assert other_project_response.status_code == 403
    assert task.sources.count() == 0


@pytest.mark.django_db(transaction=True)
def test_full_generation_review_and_export_workflow(client, tmp_path, monkeypatch):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    source_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sources/",
        {"document_version_ids": [case["source_version"].pk]},
        content_type="application/json",
    )
    queued_ids: list[str] = []
    monkeypatch.setattr(
        "apps.document_generation.queues.queue_generation_task",
        lambda task_id: queued_ids.append(task_id),
    )
    extract_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/extract/",
        content_type="application/json",
    )
    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        assert run_generation_task(str(task.pk)) == "extracted"
    task.refresh_from_db()
    assert task.status == GenerationTask.Status.NEEDS_CONFIRMATION
    assert task.facts_snapshot[0]["evidence"]
    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        confirm_response = client.put(
            f"/api/v1/document-generation/tasks/{task.pk}/facts/confirm/",
            {
                "facts": [
                    {
                        "field": "project_name",
                        "value": "风场入场项目",
                        "value_type": "string",
                        "source_document_version_id": case["source_version"].pk,
                        "locator": {
                            "paragraph_index": 0,
                            "text_quote": "项目名称：风场入场项目",
                        },
                        "confidence": 1,
                    }
                ]
            },
            content_type="application/json",
        )
    queued_ids.clear()
    generate_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/generate/",
        content_type="application/json",
    )
    task.refresh_from_db()

    assert source_response.status_code == 200
    assert extract_response.status_code == 202
    assert extract_response.json()["status"] == GenerationTask.Status.EXTRACTING
    assert confirm_response.json()["status"] == GenerationTask.Status.READY
    assert generate_response.status_code == 202
    assert task.status == GenerationTask.Status.QUEUED
    assert queued_ids == [str(task.pk)]

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        assert run_generation_task(str(task.pk)) == "completed"
        assert run_generation_task(str(task.pk)) == "skipped"
    task.refresh_from_db()
    section = GeneratedSection.objects.get(task=task, section_code="overview")
    assert task.status == GenerationTask.Status.REVIEW_REQUIRED
    assert task.generation_attempts == 2
    assert task.prompt_version == "section_generation/v4"
    assert section.content
    assert (tmp_path / task.draft_storage_path).is_file()

    regenerate_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sections/overview/regenerate/",
        {
            "instruction": "补充岗位分工：张三，12345678，并增加检查记录要求。",
            "rag_chunk_ids": [],
        },
        content_type="application/json",
    )
    task.refresh_from_db()
    regeneration_review = task.reviews.get(
        action=GenerationReview.Action.SECTION_REGENERATED,
    )
    assert regenerate_response.status_code == 202
    assert task.status == GenerationTask.Status.QUEUED
    assert regeneration_review.comment.startswith("补充岗位分工")
    assert regeneration_review.metadata["conversation_status"] == "queued"
    assert regeneration_review.metadata["required_literals"] == ["张三", "12345678"]

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        assert run_generation_task(str(task.pk)) == "completed"
    task.refresh_from_db()
    section.refresh_from_db()
    regeneration_review.refresh_from_db()
    assert task.status == GenerationTask.Status.REVIEW_REQUIRED
    assert section.revision == 2
    assert "张三" in section.content
    assert "12345678" in section.content
    assert regeneration_review.metadata["conversation_status"] == "completed"
    assert regeneration_review.metadata["revision_after"] == 2

    followup_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sections/overview/regenerate/",
        {
            "instruction": "没有看到你在正文内容中的修改。",
            "rag_chunk_ids": [],
        },
        content_type="application/json",
    )
    assert followup_response.status_code == 202
    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        assert run_generation_task(str(task.pk)) == "review_required"
    task.refresh_from_db()
    section.refresh_from_db()
    followup_review = task.reviews.filter(
        action=GenerationReview.Action.SECTION_REGENERATED,
    ).latest("id")
    assert task.status == GenerationTask.Status.REVIEW_REQUIRED
    assert section.revision == 2
    assert "张三" in section.content
    assert followup_review.metadata["conversation_status"] == "failed"
    assert "未能落实到正文" in followup_review.metadata["assistant_message"]

    lock_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/sections/overview/lock/",
        {"locked": True},
        content_type="application/json",
    )
    submit_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/submit-review/",
        {"comment": "编制人员已复核"},
        content_type="application/json",
    )
    approve_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/approve/",
        {"comment": "技术负责人确认"},
        content_type="application/json",
    )
    export_info_before = client.get(
        f"/api/v1/document-generation/tasks/{task.pk}/export-info/",
    )
    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        export_response = client.post(
            f"/api/v1/document-generation/tasks/{task.pk}/export/",
            {
                "idempotency_key": "export-001",
                "filename": "风场项目-四措两案-终版.docx",
            },
            content_type="application/json",
        )
        repeated_export = client.post(
            f"/api/v1/document-generation/tasks/{task.pk}/export/",
            {"idempotency_key": "export-001"},
            content_type="application/json",
        )
    export_info_after = client.get(
        f"/api/v1/document-generation/tasks/{task.pk}/export-info/",
    )
    task.refresh_from_db()
    output = task.output_document_version.document

    assert lock_response.status_code == 200
    assert submit_response.status_code == 200
    assert submit_response.json()["status"] == GenerationTask.Status.PENDING_APPROVAL
    assert approve_response.status_code == 200
    assert export_info_before.status_code == 200
    assert export_info_before.json()["target_folder"] == "技术方案"
    assert export_info_before.json()["agent_generated_count"] == 0
    assert export_info_before.json()["default_filename"].endswith(".docx")
    assert export_response.status_code == 200
    assert repeated_export.status_code == 200
    assert export_info_after.status_code == 200
    assert export_info_after.json()["agent_generated_count"] == 1
    assert task.status == GenerationTask.Status.EXPORTED
    assert output.project == case["project"]
    assert output.folder == case["technical_folder"]
    assert output.title == "风场项目-四措两案-终版"
    assert output.current_version.original_filename == "风场项目-四措两案-终版.docx"
    assert "不是检测报告或完工报告" in output.description
    assert (
        GenerationReview.objects.filter(
            task=task,
            action=GenerationReview.Action.EXPORTED,
        ).count()
        == 1
    )
    events_response = client.get(
        f"/api/v1/document-generation/tasks/{task.pk}/events/",
        {"after_sequence": 0},
    )
    assert events_response.status_code == 200
    assert any(event["event_type"] == "rag" for event in events_response.json())
    assert any(event["tool"] == "render_word_document" for event in events_response.json())


@pytest.mark.django_db
def test_export_name_conflict_reports_existing_agent_file_count(client, tmp_path):
    case = setup_generation_case(tmp_path)
    client.force_login(case["manager"])
    task = create_task_via_api(client, case)
    task.status = GenerationTask.Status.APPROVED
    task.save(update_fields=["status", "updated_at"])
    existing_version = create_stored_version(
        root=tmp_path,
        actor=case["manager"],
        project=case["project"],
        folder=case["technical_folder"],
        title="已有四措两案",
        filename="已有四措两案.docx",
        content=docx_bytes("已导出的Agent文件"),
    )
    GenerationTask.objects.create(
        project=case["project"],
        template=case["template"],
        status=GenerationTask.Status.EXPORTED,
        progress=100,
        idempotency_key="previous-agent-export",
        request_fingerprint="previous-agent-export",
        output_document_version=existing_version,
        export_idempotency_key="previous-export-key",
        created_by=case["manager"],
    )

    info_response = client.get(
        f"/api/v1/document-generation/tasks/{task.pk}/export-info/",
    )
    export_response = client.post(
        f"/api/v1/document-generation/tasks/{task.pk}/export/",
        {
            "idempotency_key": "conflicting-export",
            "filename": "已有四措两案.docx",
        },
        content_type="application/json",
    )

    assert info_response.status_code == 200
    assert info_response.json()["agent_generated_count"] == 1
    assert export_response.status_code == 409
    assert "已有 1 份 Agent 生成文件" in export_response.json()["message"]
    assert "请修改文件名后重试" in export_response.json()["message"]


@pytest.mark.django_db
@override_settings(DOCUMENT_AGENT_ENABLED=False)
def test_feature_flag_off_leaves_api_unavailable(client):
    admin = make_user("admin", User.Role.SYSTEM_ADMIN)
    client.force_login(admin)

    response = client.get("/api/v1/document-generation/templates/")

    assert response.status_code == 404
    assert response.json()["code"] == "DOCUMENT_AGENT_DISABLED"


@pytest.mark.django_db
@override_settings(DOCUMENT_AGENT_ENABLED=True, DOCUMENT_AGENT_PHASE5_APPROVED=False)
def test_phase5_gate_blocks_accidental_api_activation(client):
    admin = make_user("admin", User.Role.SYSTEM_ADMIN)
    client.force_login(admin)

    response = client.get("/api/v1/document-generation/templates/")

    assert response.status_code == 503
    assert response.json()["code"] == "PHASE5_GATE_NOT_APPROVED"
