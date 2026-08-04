from __future__ import annotations

from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from docx import Document as DocxDocument

from apps.documents.models import Document
from apps.documents.services import create_document
from apps.folders.models import Folder

from ..jobs import run_knowledge_corpus_upload
from ..models import ApprovalStatus, KnowledgeCorpusUpload, KnowledgeSection
from ..technical_solution_corpus import (
    enqueue_technical_solution_corpus,
    scan_technical_solution_corpus,
)

User = get_user_model()
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("安全措施", level=1)
    document.add_paragraph(
        "进入风机前执行停机、验电和挂牌上锁，高处作业人员必须正确佩戴全身式安全带。"
    )
    document.add_paragraph("相控阵超声检测作业区域设置警戒线，雨雪、大风等不利天气停止登塔作业。")
    document.add_heading("应急预案", level=1)
    document.add_paragraph("发生人员受伤时立即停止作业，启动现场救援并按程序上报。")
    document.add_heading("环境保护", level=1)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _entry_material_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("入场任务通知")
    document.add_paragraph("本次作业范围为风机塔筒及高强度螺栓检测。")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _user(username: str, role: str):
    return User.objects.create_user(
        username=username,
        password="Password123!",
        real_name=username,
        role=role,
    )


def _public_technical_folder(admin):
    return Folder.objects.create(
        project=None,
        parent=None,
        name="技术方案",
        code="PUBLIC-TECH-SOLUTION",
        is_system_root=True,
        created_by=admin,
    )


@pytest.mark.django_db(transaction=True)
@override_settings(
    DOCUMENT_AGENT_ENABLED=True,
    DOCUMENT_AGENT_PHASE5_APPROVED=True,
    DOCUMENT_AGENT_ALLOW_FAKE_PROVIDER=True,
)
def test_admin_upload_is_queued_and_embedded_into_approved_knowledge(
    client,
    tmp_path,
    monkeypatch,
):
    admin = _user("corpus-admin", User.Role.SYSTEM_ADMIN)
    _public_technical_folder(admin)
    queued: list[str] = []
    monkeypatch.setattr(
        "apps.document_generation.knowledge_corpus.queue_knowledge_corpus_upload",
        lambda upload_id: queued.append(upload_id),
    )
    client.force_login(admin)

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.post(
            "/api/v1/document-generation/knowledge-uploads/",
            {
                "file": SimpleUploadedFile(
                    "safety-measures.docx",
                    _docx_bytes(),
                    content_type=DOCX_MIME,
                ),
                "section_codes": [
                    "safety_measures",
                    "emergency_plan",
                    "environmental_measures",
                ],
            },
        )

        assert response.status_code == 202
        upload = KnowledgeCorpusUpload.objects.get(pk=response.json()["id"])
        assert upload.status == KnowledgeCorpusUpload.Status.QUEUED
        assert upload.section_codes == [
            "safety_measures",
            "emergency_plan",
            "environmental_measures",
        ]
        assert upload.fallback_to_full_document is False
        assert (
            upload.source_document_version.document.access_level == Document.AccessLevel.RESTRICTED
        )
        assert queued == [str(upload.pk)]

        assert run_knowledge_corpus_upload(str(upload.pk)) == "completed"

    upload.refresh_from_db()
    assert upload.status == KnowledgeCorpusUpload.Status.SUCCEEDED
    assert upload.chunk_count > 0
    assert upload.embedding_model_alias == "hashing-fake-v1"
    assert upload.embedding_dimension == 64
    assert upload.indexed_section_codes == ["safety_measures", "emergency_plan"]
    assert upload.skipped_section_codes == ["environmental_measures"]
    chunks = KnowledgeSection.objects.filter(source_document_version=upload.source_document_version)
    assert chunks.count() == upload.chunk_count
    assert set(chunks.values_list("section_code", flat=True)) == {
        "safety_measures",
        "emergency_plan",
    }
    assert (
        chunks.filter(
            approval_status=ApprovalStatus.APPROVED,
            is_active=True,
            approved_by=admin,
            embedding_model_alias="hashing-fake-v1",
            embedding_dimension=64,
        ).count()
        == upload.chunk_count
    )


@pytest.mark.django_db(transaction=True)
@override_settings(
    DOCUMENT_AGENT_ENABLED=True,
    DOCUMENT_AGENT_PHASE5_APPROVED=True,
    DOCUMENT_AGENT_ALLOW_FAKE_PROVIDER=True,
)
def test_single_section_entry_material_falls_back_to_whole_document(
    client,
    tmp_path,
    monkeypatch,
):
    admin = _user("entry-corpus-admin", User.Role.SYSTEM_ADMIN)
    _public_technical_folder(admin)
    monkeypatch.setattr(
        "apps.document_generation.knowledge_corpus.queue_knowledge_corpus_upload",
        lambda upload_id: None,
    )
    client.force_login(admin)

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        response = client.post(
            "/api/v1/document-generation/knowledge-uploads/",
            {
                "file": SimpleUploadedFile(
                    "entry-notice.docx",
                    _entry_material_bytes(),
                    content_type=DOCX_MIME,
                ),
                "section_codes": ["overview"],
            },
        )
        assert response.status_code == 202
        upload = KnowledgeCorpusUpload.objects.get(pk=response.json()["id"])
        assert upload.fallback_to_full_document is True
        assert run_knowledge_corpus_upload(str(upload.pk)) == "completed"

    upload.refresh_from_db()
    assert upload.indexed_section_codes == ["overview"]
    assert upload.skipped_section_codes == []
    assert KnowledgeSection.objects.filter(
        source_document_version=upload.source_document_version,
        section_code="overview",
    ).exists()


@pytest.mark.django_db(transaction=True)
@override_settings(
    DOCUMENT_AGENT_ENABLED=True,
    DOCUMENT_AGENT_PHASE5_APPROVED=True,
    DOCUMENT_AGENT_ALLOW_FAKE_PROVIDER=True,
)
def test_technical_solution_scan_queues_only_usable_missing_sections(
    tmp_path,
    monkeypatch,
):
    admin = _user("bulk-corpus-admin", User.Role.SYSTEM_ADMIN)
    folder = _public_technical_folder(admin)
    queued: list[str] = []
    monkeypatch.setattr(
        "apps.document_generation.technical_solution_corpus.queue_knowledge_corpus_upload",
        lambda upload_id: queued.append(upload_id),
    )

    with override_settings(FILE_STORAGE_ROOT=tmp_path):
        document = create_document(
            actor=admin,
            folder=folder,
            uploaded_file=SimpleUploadedFile(
                "complete-plan.docx",
                _docx_bytes(),
                content_type=DOCX_MIME,
            ),
            title="complete-plan.docx",
        )
        plans = scan_technical_solution_corpus()
        plan = next(item for item in plans if item.document_id == document.pk)
        assert plan.section_codes == (
            "safety_measures",
            "emergency_plan",
            "environmental_measures",
        )
        assert plan.empty_section_codes == ("environmental_measures",)
        assert plan.estimated_chunk_count > 0

        uploads = enqueue_technical_solution_corpus(actor=admin, plans=plans)
        assert len(uploads) == 1
        assert queued == [str(uploads[0].pk)]
        assert run_knowledge_corpus_upload(str(uploads[0].pk)) == "completed"

    uploads[0].refresh_from_db()
    assert uploads[0].indexed_section_codes == [
        "safety_measures",
        "emergency_plan",
    ]
    assert uploads[0].skipped_section_codes == ["environmental_measures"]


@pytest.mark.django_db
@override_settings(
    DOCUMENT_AGENT_ENABLED=True,
    DOCUMENT_AGENT_PHASE5_APPROVED=True,
)
def test_non_admin_cannot_upload_rag_corpus(client):
    operator = _user("corpus-operator", User.Role.DATA_OPERATOR)
    client.force_login(operator)

    response = client.post(
        "/api/v1/document-generation/knowledge-uploads/",
        {
            "file": SimpleUploadedFile(
                "safety-measures.docx",
                _docx_bytes(),
                content_type=DOCX_MIME,
            ),
            "section_codes": ["safety_measures"],
        },
    )

    assert response.status_code == 403
